"""
Analyse generique de tables (CSV/Excel) : fonctionne quelles que soient les
colonnes reellement presentes, sans dependre de la structure exacte du
dictionnaire. Detecte automatiquement les colonnes d'identifiant et de date
par le nom des colonnes / le contenu.

Regle de securite : toute colonne qui ressemble a un nom/prenom est
systematiquement exclue des resultats affiches (meme si elle est presente
dans le fichier source), conformement a la regle de non-exposition des
donnees nominatives de l'assistant OPO.
"""

from __future__ import annotations

import io
import re
import unicodedata
from datetime import datetime

import numpy as np
import pandas as pd


def _sans_accents(texte: str) -> str:
    """Retire les accents d'un texte (meme principe que app.sans_accents,
    duplique ici pour que data_tools.py reste importable independamment de
    app.py). Necessaire pour reconnaitre le nom ou l'alias d'une table tape
    avec ou sans accent (ex: "décès"/"deces")."""
    return "".join(c for c in unicodedata.normalize("NFD", str(texte)) if unicodedata.category(c) != "Mn")

NAME_LIKE = re.compile(r"(nom|name|prenom|prénom|surname|firstname|lastname)", re.IGNORECASE)
ID_LIKE = re.compile(r"(id|identif)", re.IGNORECASE)
# Volontairement restreint a des motifs qui identifient vraiment l'AGENT (pas
# n'importe quelle colonne liee a une enquete) : le motif large precedent
# (simple "enquet"/"interview" en sous-chaine) capturait a tort des colonnes
# comme `enquete_id` ou `date_enquete`, presentes sur la quasi-totalite des
# tables du schema reel de l'observatoire mais qui ne designent PAS l'agent -
# seulement l'evenement d'enquete/visite auquel la fiche se rattache.
AGENT_LIKE = re.compile(
    r"(field_?wrkr|fieldworker|agent_?id|agent_?name|^agent$|agent_?enquet|interviewer_?id)", re.IGNORECASE
)
# `agent_?name` ajoute suite a l'inspection des vrais fichiers exportes par
# l'observatoire : la quasi-totalite des tables d'evenements reelles
# (opo_hypervel_arrivees/deces/naissances/education/...) portent une colonne
# `agent_name` qui contient DEJA le nom complet reel de l'agent enqueteur
# (verifie sur `opo_hypervel_education.csv` : memes noms que dans le fichier
# equipe.dta fourni par l'observatoire, ex "BADINI RACHIDE") - la reponse la
# plus directe possible au besoin "je veux que les noms des agents
# apparaissent", sans meme avoir besoin d'une jointure vers une autre table.

# Identifiants REELS documentes dans le schema relationnel de l'observatoire
# (voir data/docs/00_schema_relations.txt, construit a partir du dictionnaire
# de donnees et du document de correspondance des tables, PLUS confirmes par
# inspection directe des 28 vraies tables exportees - jamais invente ni
# devine par convention de nommage). Consigne explicite de l'observatoire :
# un nom de colonne qui RESSEMBLE a un identifiant (ex: "menage_id",
# "individu_id", "uch_id", "enquete_id" - la convention Laravel/Hypervel
# `<table>_id`, presente EN PARALLELE des vrais identifiants sur la quasi
# totalite des tables reelles) n'en est pas la preuve : le menage est
# identifie par `socialgpid`, jamais par `menage_id`. Seule cette liste, et
# la documentation dont elle vient, fait foi.
IDENTIFIANTS_REELS_DOCUMENTES = [
    # --- Cles des 3 entites centrales + episodes/evenements (voir
    # data/docs/00_schema_relations.txt, synthese officielle des relations
    # entre les 23 tables du schema HDSS/ODSS d'origine). ---
    "individid", "socialgpid", "locationid", "episodeid", "episodeid_res",
    "episodeid_head", "eventid", "observeid", "sobserveid", "eobserveid",
    "fatherid", "motherid", "headid", "individid2", "childid", "ownerid",
    "owner_id", "pregoutid",
    # --- Confirmes en inspectant directement les 28 vraies tables
    # opo_hypervel_* exportees par l'observatoire (colonnes reellement
    # presentes, pas seulement documentees) : respondid sur la quasi-totalite
    # des tables, srespondid/erespondid sur les tables d'episodes (meme
    # convention documentee que sobserveid/eobserveid), peventid present tel
    # quel dans opo_hypervel_issue_grossesses.csv. ---
    "respondid", "srespondid", "erespondid", "peventid",
    # --- Consigne explicite de l'observatoire (2026-08-11) : reprendre TOUS
    # les identifiants marques "Identifiant" dans le dictionnaire de
    # donnees source, pas seulement le sous-ensemble deja verifie present
    # dans les tables actuellement chargees - une table/un lot d'export
    # different peut porter des colonnes non vues jusqu'ici. Liste ci-dessous
    # extraite de "source_manuel technique et d'utilisation des donnees du
    # ODSS (juil04).txt" (colonne "Identifiant" du dictionnaire technique
    # d'origine), en plus de celles deja listees ci-dessus. Aucune n'a ete
    # devinee par convention de nommage : chacune est explicitement typee
    # "Identifiant" dans le document source. ---
    "accomid", "chiefid", "socialgpidtmp", "lobserveid", "slobserveid",
    "elobserveid", "region", "cluster",
    "shusb_region", "shusb_locationid", "ehusb_region", "ehusb_locationid",
    "swife_region", "swife_locationid", "ewife_region", "ewife_locationid",
]


def load_table(path: str) -> pd.DataFrame:
    """Charge une table depuis un fichier Excel/Stata/CSV.

    Bug reel corrige ici, decouvert en inspectant les vrais fichiers exportes
    par l'observatoire : `pd.read_csv(path)` suppose une virgule comme
    separateur, mais environ la moitie des exports reels du schema
    opo_hypervel_* utilisent un POINT-VIRGULE (export Excel en locale
    francaise, ou la virgule est deja le separateur decimal). Sans detection
    du separateur, ces fichiers se chargeaient en UNE SEULE colonne (tout le
    contenu de chaque ligne concatene comme texte), ce qui pouvait meme
    aboutir a une table vide (0 colonne) selon le contenu - une table
    entiere silencieusement invisible pour l'assistant, sans aucune erreur
    affichee. `sep=None, engine="python"` laisse pandas detecter le vrai
    separateur (virgule ou point-virgule) a partir du contenu reel du
    fichier, verifie sur l'ensemble des 28 tables reelles de l'observatoire."""
    chemin = str(path).lower()
    if chemin.endswith((".xlsx", ".xls")):
        df = pd.read_excel(path)
    elif chemin.endswith(".dta"):
        df = pd.read_stata(path)
    else:
        df = pd.read_csv(path, sep=None, engine="python")
    return strip_names(df)


def est_classeur_excel(path: str) -> bool:
    return str(path).lower().endswith((".xlsx", ".xls"))


def charger_classeur(path: str) -> dict[str, pd.DataFrame]:
    """Charge TOUTES les feuilles d'un classeur Excel comme autant de tables
    distinctes (au lieu de ne lire que la premiere feuille), pour reconnaitre
    directement un classeur qui contient plusieurs tables (ex: une feuille
    par table de l'observatoire dans un seul fichier).

    Renvoie un dict {nom_de_feuille: dataframe}, en ignorant les feuilles
    vides et en retirant les colonnes nominatives de chaque feuille (memes
    garde-fous que load_table)."""
    classeur = pd.ExcelFile(path)
    tables = {}
    for nom_feuille in classeur.sheet_names:
        df = classeur.parse(nom_feuille)
        if df.empty or len(df.columns) == 0:
            continue
        tables[nom_feuille] = strip_names(df)
    return tables


def strip_names(df: pd.DataFrame) -> pd.DataFrame:
    """Retire toute colonne qui ressemble a un nom/prenom d'une PERSONNE
    ENQUETEE avant analyse (vie privee des repondants/individus) - a
    l'exception du nom de l'AGENT enqueteur (ex: `agent_name`, reconnu par
    `AGENT_LIKE`), qui n'est pas une donnee de vie privee des personnes
    enquetees mais une donnee operationnelle sur le personnel de terrain,
    necessaire au suivi de performance par agent explicitement demande par
    l'observatoire (voir `rapport_agents`/`fusion_agent_controleur`). Bug
    reel corrige ici : `agent_name` contient "name" et etait donc
    silencieusement supprimee comme n'importe quel nom de repondant, rendant
    impossible l'affichage du nom de l'agent alors que la donnee existe deja
    dans les vraies tables de l'observatoire."""
    to_drop = [c for c in df.columns if NAME_LIKE.search(str(c)) and not AGENT_LIKE.search(str(c))]
    if to_drop:
        df = df.drop(columns=to_drop)
    return df


def detect_id_columns(df: pd.DataFrame) -> list[str]:
    return [c for c in df.columns if ID_LIKE.search(str(c))]


def detect_agent_columns(df: pd.DataFrame) -> list[str]:
    """Colonnes identifiant l'agent enqueteur/le personnel de terrain ayant
    saisi une fiche (ex: `field_wrkr`), detectees par leur nom - pour un
    controle de qualite/performance par agent sans configuration manuelle."""
    return [c for c in df.columns if AGENT_LIKE.search(str(c))]


def detect_date_columns(df: pd.DataFrame) -> list[str]:
    date_cols = []
    for c in df.columns:
        if pd.api.types.is_datetime64_any_dtype(df[c]):
            date_cols.append(c)
            continue
        if df[c].dtype == object:
            sample = df[c].dropna().astype(str).head(20)
            if len(sample) == 0:
                continue
            parsed = pd.to_datetime(sample, errors="coerce", dayfirst=True)
            if parsed.notna().mean() > 0.7:
                date_cols.append(c)
    return date_cols


def repartition(df: pd.DataFrame, colonne: str) -> pd.DataFrame:
    if colonne not in df.columns:
        raise ValueError(f"Colonne '{colonne}' absente. Colonnes disponibles : {list(df.columns)}")
    out = df[colonne].value_counts(dropna=False).rename("effectif").to_frame()
    out["pourcentage"] = (100 * out["effectif"] / out["effectif"].sum()).round(1)
    return out


def echantillon(df: pd.DataFrame, n: int = 100, seed: int = 20260729) -> pd.DataFrame:
    n = min(n, len(df))
    return df.sample(n=n, random_state=seed).reset_index(drop=True)


def doublons(df: pd.DataFrame, colonne: str | None = None) -> pd.DataFrame:
    """Doublons sur une colonne d'identifiant (auto-detectee si non precisee)."""
    if colonne is None:
        ids = detect_id_columns(df)
        if not ids:
            raise ValueError("Aucune colonne d'identifiant detectee automatiquement ; precisez `colonne`.")
        colonne = ids[0]
    counts = df[colonne].value_counts()
    dup_values = counts[counts > 1].index
    return df[df[colonne].isin(dup_values)].sort_values(colonne)


class RequeteInvalide(ValueError):
    """Levee quand la specification de requete (operation/colonne/valeur)
    ne correspond a rien d'exploitable dans la table chargee - jamais
    d'execution "au hasard" sur une colonne qui n'existe pas ou une
    operation non reconnue."""


def _valeur_numerique(valeur) -> float:
    try:
        return float(str(valeur).strip().replace(",", "."))
    except (TypeError, ValueError):
        return float("nan")


# Chaque operateur prend la Series de la colonne filtree et la valeur brute
# (fournie par le classifieur LLM, jamais executee comme du code) et renvoie
# un masque booleen - comparaison textuelle insensible a la casse/aux espaces
# pour "=="/"!="/"contient" (les identifiants/modalites de l'observatoire sont
# souvent des codes ou libelles), conversion numerique pour les comparaisons
# d'ordre.
OPERATEURS_REQUETE = {
    "==": lambda s, v: s.astype(str).str.strip().str.lower() == str(v).strip().lower(),
    "!=": lambda s, v: s.astype(str).str.strip().str.lower() != str(v).strip().lower(),
    ">": lambda s, v: pd.to_numeric(s, errors="coerce") > _valeur_numerique(v),
    "<": lambda s, v: pd.to_numeric(s, errors="coerce") < _valeur_numerique(v),
    ">=": lambda s, v: pd.to_numeric(s, errors="coerce") >= _valeur_numerique(v),
    "<=": lambda s, v: pd.to_numeric(s, errors="coerce") <= _valeur_numerique(v),
    "contient": lambda s, v: s.astype(str).str.contains(str(v), case=False, na=False, regex=False),
}


def executer_requete_donnees(
    df: pd.DataFrame, operation: str, colonne_cible: str | None, filtres: list[dict] | None
) -> dict:
    """Execute une requete de type filtre + agregat (compter/lister/moyenne/
    somme/min/max) sur une table REELLEMENT chargee, a partir d'une
    specification produite par `rag.classifier_intention` (action REQUETE).

    C'est ce qui permet de repondre a une question de calcul precis sur les
    donnees reelles ("combien de naissances a Ouahigouya en 2026 ?", "age
    moyen des mères dont la grossesse est en cours") sans se limiter aux
    analyses fixes (repartition/echantillon/doublons/coherence) ni retomber
    sur le dictionnaire documentaire des que la question sort de ces 4 cas.

    Toute colonne de `filtres`/`colonne_cible` qui n'existe pas reellement
    dans `df` est ignoree (filtre) ou leve `RequeteInvalide` (colonne_cible
    d'une agregation) plutot que d'echouer silencieusement sur un mauvais
    calcul - jamais de confiance aveugle dans ce qu'a produit le LLM."""
    operation = (operation or "").strip().lower()
    if operation not in ("compter", "lister", "moyenne", "somme", "min", "max"):
        raise RequeteInvalide(f"Opération « {operation} » non reconnue.")

    masque = pd.Series(True, index=df.index)
    filtres_appliques = []
    for f in filtres or []:
        col = f.get("colonne")
        op = f.get("operateur", "==")
        val = f.get("valeur")
        if col not in df.columns or op not in OPERATEURS_REQUETE:
            continue
        try:
            masque &= OPERATEURS_REQUETE[op](df[col], val)
        except Exception:
            continue
        filtres_appliques.append(f"{col} {op} {val}")

    sous_ensemble = df.loc[masque]

    if operation == "compter":
        return {"operation": operation, "resultat": int(len(sous_ensemble)), "filtres_appliques": filtres_appliques}

    if operation == "lister":
        return {
            "operation": operation, "resultat": sous_ensemble.head(50).reset_index(drop=True),
            "n_total": int(len(sous_ensemble)), "filtres_appliques": filtres_appliques,
        }

    if colonne_cible not in df.columns:
        raise RequeteInvalide(f"Colonne cible « {colonne_cible} » introuvable pour l'opération « {operation} ».")

    valeurs = pd.to_numeric(sous_ensemble[colonne_cible], errors="coerce").dropna()
    if valeurs.empty:
        return {
            "operation": operation, "resultat": None, "colonne_cible": colonne_cible,
            "filtres_appliques": filtres_appliques,
        }
    fonction = {"moyenne": "mean", "somme": "sum", "min": "min", "max": "max"}[operation]
    return {
        "operation": operation, "resultat": float(getattr(valeurs, fonction)()), "colonne_cible": colonne_cible,
        "n_valeurs": int(len(valeurs)), "filtres_appliques": filtres_appliques,
    }


class RequeteSQLInvalide(ValueError):
    """Levee quand une requete SQL (generee par le LLM, voir
    `rag.generer_requete_sql`) n'est pas une simple lecture (SELECT/WITH) ou
    echoue a l'execution - jamais d'execution aveugle d'une instruction
    potentiellement destructive ou d'une requete syntaxiquement invalide."""


# Mots-cles qui, s'ils apparaissent n'importe ou dans la requete, la rendent
# refusee d'office - une deuxieme ligne de defense en plus de l'exigence
# "commence par SELECT/WITH", au cas ou le LLM glisserait une instruction de
# modification dans une sous-requete ou un commentaire.
_SQL_INTERDIT = re.compile(
    r"\b(INSERT|UPDATE|DELETE|DROP|ALTER|CREATE|ATTACH|DETACH|COPY|PRAGMA|CALL|EXPORT|IMPORT|"
    r"INSTALL|LOAD|SET|GRANT|REVOKE|VACUUM|CHECKPOINT)\b",
    re.IGNORECASE,
)


def executer_sql(tables: dict, requete_sql: str) -> pd.DataFrame:
    """Execute une requete SQL en LECTURE SEULE sur les tables reellement
    chargees, via DuckDB (qui sait interroger directement des DataFrame
    pandas enregistres, sans copie sur disque) - c'est ce qui permet de
    repondre a une question qui doit croiser 2, 3, 4 tables ou plus a la fois
    (jointures, filtres, agregations, groupby), au-dela de ce que l'action
    REQUETE mono-table (`executer_requete_donnees`) ou le controle croise
    fixe deces/depart (`controle_deces_present`) peuvent couvrir.

    Garde-fous avant toute execution : une seule instruction, qui doit
    commencer par SELECT ou WITH, et ne doit contenir aucun mot-cle de
    modification/administration (voir `_SQL_INTERDIT`) - la requete est
    TOUJOURS generee par un LLM (jamais tapee par un utilisateur final dans
    ce flux), mais ne doit jamais etre executee sans validation.

    Leve `RequeteSQLInvalide` si la requete ne passe pas ces garde-fous ou si
    son execution echoue (colonne/table inexistante, erreur de syntaxe...) -
    jamais d'exception DuckDB brute propagee a l'appelant."""
    requete_sql = (requete_sql or "").strip().rstrip(";").strip()
    if not requete_sql:
        raise RequeteSQLInvalide("Requête vide.")
    if ";" in requete_sql:
        raise RequeteSQLInvalide("Plusieurs instructions ne sont pas autorisées en une seule requête.")
    if not re.match(r"^(SELECT|WITH)\b", requete_sql, re.IGNORECASE):
        raise RequeteSQLInvalide("Seules les requêtes SELECT (ou WITH ... SELECT) sont autorisées.")
    if _SQL_INTERDIT.search(requete_sql):
        raise RequeteSQLInvalide("Instruction non autorisée détectée dans la requête.")

    import duckdb  # importe seulement ici : evite le cout de chargement pour
    # tout le reste du module quand cette fonctionnalite n'est pas utilisee.

    connexion = duckdb.connect(database=":memory:")
    try:
        for nom, df in tables.items():
            connexion.register(nom, df)
        resultat = connexion.execute(requete_sql).fetchdf()
    except RequeteSQLInvalide:
        raise
    except Exception as e:
        raise RequeteSQLInvalide(f"Erreur d'exécution SQL : {e}") from e
    finally:
        connexion.close()
    return resultat.head(200).reset_index(drop=True)


def dates_incoherentes(df: pd.DataFrame, colonne: str, borne_min: str = "1900-01-01") -> pd.DataFrame:
    parsed = pd.to_datetime(df[colonne], errors="coerce", dayfirst=True)
    aujourdhui = pd.Timestamp(datetime.now())
    mask = parsed.isna() | (parsed > aujourdhui) | (parsed < pd.Timestamp(borne_min))
    return df[mask]


def rapport_coherence(df: pd.DataFrame) -> dict:
    """Rapport generique : doublons sur les colonnes d'ID detectees, dates
    invraisemblables sur les colonnes de date detectees. Aucune correction
    n'est appliquee : uniquement un signalement, a valider par une personne
    habilitee.

    Le rapport indique explicitement quelles colonnes ont ete examinees
    (`colonnes_id_verifiees`, `colonnes_date_verifiees`), pour que la reponse
    reste precise sur ce qui a ete effectivement controle - plutot qu'un
    simple "aucune anomalie" qui pourrait laisser croire a un controle
    exhaustif alors qu'aucune colonne pertinente n'aurait ete detectee."""
    colonnes_id = detect_id_columns(df)
    colonnes_date = detect_date_columns(df)

    rapport = {
        "n_lignes": len(df),
        "colonnes": list(df.columns),
        "colonnes_id_verifiees": colonnes_id,
        "colonnes_date_verifiees": colonnes_date,
        "anomalies": {},
    }

    for id_col in colonnes_id:
        dups = doublons(df, id_col)
        if len(dups) > 0:
            rapport["anomalies"][f"doublons::{id_col}"] = len(dups)

    for date_col in colonnes_date:
        bad = dates_incoherentes(df, date_col)
        if len(bad) > 0:
            rapport["anomalies"][f"dates_invraisemblables::{date_col}"] = len(bad)

    taux_manquants = df.isna().mean().round(3)
    rapport["taux_valeurs_manquantes"] = taux_manquants[taux_manquants > 0].to_dict()

    return rapport


def rapport_agents(df: pd.DataFrame, colonne_agent: str | None = None) -> pd.DataFrame:
    """Rapport de performance/qualite par agent enqueteur : nombre de fiches
    saisies et indicateurs d'anomalies (doublons d'identifiant, dates
    invraisemblables, taux de valeurs manquantes) par agent - pour reperer
    une charge de travail inhabituelle ou un agent avec beaucoup plus
    d'erreurs que les autres (controle qualite du travail de terrain).

    Si `colonne_agent` n'est pas precisee, la premiere colonne d'agent
    detectee automatiquement (voir `detect_agent_columns`) est utilisee.
    Leve une ValueError si aucune colonne d'agent n'est detectee ni precisee,
    ou si la colonne demandee n'existe pas."""
    if colonne_agent is None:
        agents = detect_agent_columns(df)
        if not agents:
            raise ValueError(
                "Aucune colonne d'agent enquêteur détectée automatiquement (ex: field_wrkr) ; "
                "précisez `colonne_agent`."
            )
        colonne_agent = agents[0]
    elif colonne_agent not in df.columns:
        raise ValueError(f"Colonne '{colonne_agent}' absente. Colonnes disponibles : {list(df.columns)}")

    colonnes_id = [c for c in detect_id_columns(df) if c != colonne_agent]
    colonnes_date = detect_date_columns(df)

    masques_doublons = {c: df[c].duplicated(keep=False) for c in colonnes_id}
    masques_dates = {}
    aujourdhui = pd.Timestamp(datetime.now())
    for c in colonnes_date:
        parsed = pd.to_datetime(df[c], errors="coerce", dayfirst=True)
        masques_dates[c] = parsed.isna() | (parsed > aujourdhui) | (parsed < pd.Timestamp("1900-01-01"))

    lignes = []
    for agent, index_groupe in df.groupby(colonne_agent, dropna=False).groups.items():
        n_fiches = len(index_groupe)
        n_doublons = sum(int(masques_doublons[c].loc[index_groupe].sum()) for c in colonnes_id)
        n_dates = sum(int(masques_dates[c].loc[index_groupe].sum()) for c in colonnes_date)
        taux_manquants = float(df.loc[index_groupe].isna().mean().mean()) if n_fiches else 0.0
        lignes.append({
            "agent": agent,
            "n_fiches": n_fiches,
            "doublons_id": n_doublons,
            "dates_invraisemblables": n_dates,
            "taux_valeurs_manquantes_moyen": round(taux_manquants, 3),
        })

    return pd.DataFrame(lignes).sort_values("n_fiches", ascending=False).reset_index(drop=True)


def syntaxe_rapport_agents(nom_table: str, colonne_agent: str) -> str:
    """Syntaxe R et Stata equivalente au rapport de performance par agent
    (nombre de fiches par agent, base pour reperer les anomalies)."""
    r = f"resultat <- dplyr::count({nom_table}, {colonne_agent}, name = \"n_fiches\")"
    stata = f"use {nom_table}, clear\nbysort {colonne_agent}: gen n_fiches = _N\ntab {colonne_agent}"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def tables_mentionnees_dans_historique(historique: list[dict] | None, tables: dict) -> list[str]:
    """Recherche, dans les derniers echanges de la conversation (du plus
    recent au plus ancien), les noms de tables chargees qui ont deja ete
    mentionnes - pour qu'une question de suivi qui ne renomme pas
    explicitement la table ("et les doublons ?" apres avoir parle de
    Tindividual) reste rattachee au bon contexte plutot que de retomber sur
    un choix par defaut potentiellement different.

    `historique` est une liste de {"role": ..., "contenu": ...} (voir
    app.py:historique_recent), du plus ancien au plus recent."""
    if not historique:
        return []
    trouvees: list[str] = []
    for message in reversed(historique):
        contenu = str(message.get("contenu", "")).lower()
        for nom in tables:
            if nom.lower() in contenu and nom not in trouvees:
                trouvees.append(nom)
    return trouvees


PREFIXES_TABLE_IGNORES = (
    "fnew_", "fnew", "f_new_", "f_new",
    # Prefixe technique des exports reels de l'observatoire (base Hypervel) :
    # les tables s'appellent par ex. "opo_hypervel_individus",
    # "opo_hypervel_naissances"... - sans ce prefixe reconnu, aucun alias
    # informel ("individus", "naissances"...) ne pouvait etre reconnu, puisque
    # le nom charge ne commencait par aucun des prefixes connus jusqu'ici.
    "opo_hypervel_", "opo_hypervel",
)


def alias_table(nom: str) -> list[str]:
    """Genere des variantes informelles d'un nom de table charge, pour
    reconnaitre une mention qui omet le prefixe technique et/ou le
    singulier/pluriel exact (ex: l'equipe dit "la table education" ou
    "presence" a l'oral, alors que la table chargee s'appelle reellement
    "FNewEducation" ou "FNewPresences" - la correspondance exacte de
    sous-chaine ne suffit alors jamais a la reconnaitre).

    Ne renvoie que des alias d'au moins 4 caracteres, pour eviter qu'un
    prefixe retire ne laisse un fragment trop court et trop generique
    (risque de faux positifs sur un mot sans rapport)."""
    base = nom.lower()
    alias = base
    for prefixe in PREFIXES_TABLE_IGNORES:
        if base.startswith(prefixe) and len(base) > len(prefixe):
            alias = base[len(prefixe):]
            break

    variantes = {base, alias}
    if alias.endswith("s"):
        variantes.add(alias[:-1])
    else:
        variantes.add(alias + "s")

    # Certains noms de table reels epellent un sigle avec un underscore entre
    # chaque lettre (ex: "d_e_c_e_s" pour "deces", "c_p_n_s" pour "cpns",
    # schema Hypervel de l'observatoire) : sans underscores, l'alias redevient
    # un mot naturel que l'equipe reconnait et emploie a l'oral.
    sans_underscore = alias.replace("_", "")
    if sans_underscore != alias:
        variantes.add(sans_underscore)
        if sans_underscore.endswith("s"):
            variantes.add(sans_underscore[:-1])
        else:
            variantes.add(sans_underscore + "s")

    # Variante insensible aux accents (ex: table "Présences" -> alias
    # reconnu meme tape "presences" sans accent).
    variantes |= {_sans_accents(v) for v in list(variantes)}

    return [v for v in variantes if len(v) >= 4]


def _alias_mentionne(alias: str, q: str) -> bool:
    """Verifie qu'un alias de table (voir `alias_table`) est mentionne comme
    mot entier dans la question, plutot qu'une simple sous-chaine - evite un
    faux positif quand l'alias d'une table correspond par hasard au debut
    d'un nom de colonne sans rapport (ex: l'alias "education" de la table
    FNewEducation ne doit PAS matcher a l'interieur du nom de colonne
    "education_level", qui parle d'une colonne, pas de la table)."""
    return re.search(r"\b" + re.escape(alias) + r"\b", q) is not None


# Alias tres generiques qui coincident avec le vocabulaire courant de
# l'observatoire, au-dela du nom d'une table precise : toute la base de
# donnees parle "d'individus", "de menages", "de personnes"... Bug reel
# corrige ici - la question "les individus dont on a fait l'education ne
# sont pas dans la fiche presence" mentionne "individus" au sens general
# (les gens), pas la table opo_hypervel_individus en particulier, mais une
# correspondance nue sur ce mot la faisait quand meme compter comme "table
# mentionnee". Resultat : `resoudre_paire_tables` recevait 3 tables
# (individus, education, presences) et retenait les deux premieres au sens
# de l'ordre de chargement (education + individus), ecartant la table
# reellement visee par la question (presences). Pour ces alias-la
# seulement, une mention nue ne suffit plus : il faut un ancrage explicite
# ("dans les individus", "table individus", "base individus"...) a
# proximite immediate, comme pour n'importe quel autre nom de table cite
# explicitement.
ALIAS_AMBIGUS_GENERIQUES = {"individus", "individu", "menages", "menage", "personnes", "personne"}

_MOTS_ANCRAGE_TABLE = ("dans", "table", "base", "fiche", "feuille", "onglet")


def _alias_ancre(alias: str, q: str) -> bool:
    """Verifie qu'un alias AMBIGU (voir `ALIAS_AMBIGUS_GENERIQUES`) est bien
    introduit par un mot comme "dans"/"table"/"base"/"fiche" juste avant, et
    pas seulement employe comme mot du langage courant."""
    motif = (
        r"\b(?:" + "|".join(_MOTS_ANCRAGE_TABLE) + r")\s+"
        r"(?:la\s+|le\s+|les\s+|l['’]\s*|de\s+la\s+|des\s+)?"
        + re.escape(alias) + r"\b"
    )
    return re.search(motif, q) is not None


def resoudre_table_ciblee(
    question: str, tables: dict, nom_par_defaut: str | None = None, historique: list[dict] | None = None
):
    """Determine quelle table (parmi plusieurs deposees) est visee par une
    question, pour que toutes les tables chargees soient aussi faciles a
    interroger les unes que les autres (pas seulement la table "active") :

    1. Si le nom d'une table (ou un alias informel - voir `alias_table`) est
       mentionne explicitement dans la question, elle est prioritaire.
    2. Sinon, si un nom de colonne mentionne dans la question n'appartient
       qu'a une seule des tables chargees, cette table est retenue
       automatiquement (ex: "repartition de sex" cible directement la table
       qui contient la colonne "sex", sans avoir a la nommer ni a la
       selectionner comme table active).
    3. Sinon, si une table a ete mentionnee dans les echanges precedents de
       la conversation, elle est retenue (memoire conversationnelle - une
       question de suivi comme "et les doublons ?" reste rattachee au bon
       contexte).
    4. En dernier recours (aucun historique exploitable), on retombe sur la
       table active par defaut choisie dans l'interface.

    Renvoie un tuple (nom_de_la_table, dataframe) ou (None, None) si aucune
    table n'est disponible.
    """
    q = _sans_accents(question.lower())

    for nom in tables:
        if _sans_accents(nom.lower()) in q:
            return nom, tables[nom]

    for nom in tables:
        if any(_alias_mentionne(alias, q) for alias in alias_table(nom)):
            return nom, tables[nom]

    tables_correspondantes = []
    for nom, df in tables.items():
        if any(str(colonne).lower() in q for colonne in df.columns):
            tables_correspondantes.append(nom)
    if len(tables_correspondantes) == 1:
        nom = tables_correspondantes[0]
        return nom, tables[nom]

    for nom in tables_mentionnees_dans_historique(historique, tables):
        return nom, tables[nom]

    if nom_par_defaut is not None and nom_par_defaut in tables:
        return nom_par_defaut, tables[nom_par_defaut]
    return None, None


def tables_avec_colonne(question: str, tables: dict) -> list[str]:
    """Renvoie TOUTES les tables chargees qui contiennent une colonne
    mentionnee dans la question (pas seulement une premiere correspondance
    unique) - pour qu'une question ambigue sur une colonne presente dans
    plusieurs tables a la fois (ex: `individid` present dans 20 tables) soit
    traitee sur CHACUNE d'entre elles plutot que de silencieusement retomber
    sur un choix par defaut ("il faut tout lire, pas une seule base par
    defaut")."""
    q = question.lower()
    return [nom for nom, df in tables.items() if any(str(colonne).lower() in q for colonne in df.columns)]


def colonnes_mentionnees(question: str, tables: dict) -> list[str]:
    """Renvoie, dans l'ordre de premiere apparition dans la question, les
    noms de colonnes mentionnes - a partir de l'UNION de toutes les colonnes
    de TOUTES les tables chargees (pas seulement celles d'une table deja
    choisie par defaut), pour reconnaitre une colonne peu importe quelle
    table est regardee en premier. Sert de base a la recherche multi-table
    pour les analyses a plusieurs colonnes (bivariee, correlation,
    multivariee) quand aucune table n'est nommee explicitement."""
    q = question.lower()
    toutes_colonnes, vues = [], set()
    for df in tables.values():
        for c in df.columns:
            cle = str(c).lower()
            if cle not in vues:
                vues.add(cle)
                toutes_colonnes.append(str(c))
    return [c for c in toutes_colonnes if c.lower() in q]


def tables_avec_toutes_colonnes(colonnes: list[str], tables: dict) -> list[str]:
    """Renvoie les tables chargees qui contiennent TOUTES les colonnes
    donnees a la fois - candidates pour une analyse bivariee/multivariee sur
    plusieurs colonnes mentionnees dans la question sans qu'aucune table ne
    soit nommee explicitement ("il ne faut pas lire une seule base par
    defaut ; il faut tout lire")."""
    if not colonnes:
        return []
    colonnes_norm = {c.lower() for c in colonnes}
    resultat = []
    for nom, df in tables.items():
        colonnes_table = {str(c).lower() for c in df.columns}
        if colonnes_norm.issubset(colonnes_table):
            resultat.append(nom)
    return resultat


def resume_tables_chargees(tables: dict) -> str:
    """Decrit les tables actuellement chargees (nom, nombre de lignes et de
    colonnes) - reponse a une question "meta" sur la session en cours (ex:
    "combien de tables/feuilles sont chargees ?", "quelles tables sont
    disponibles ?"), qui ne porte pas sur le contenu d'une table mais sur ce
    qui est effectivement charge a l'instant. Ce n'est pas une question
    documentaire (le dictionnaire ne sait rien de la session en cours), donc
    elle doit etre traitee ici plutot que par le RAG."""
    if not tables:
        return "Aucune table n'est chargée pour l'instant."

    lignes = [f"**{len(tables)} table(s) chargée(s)** :"]
    for nom, df in tables.items():
        apercu_colonnes = ", ".join(f"`{c}`" for c in df.columns[:8])
        if len(df.columns) > 8:
            apercu_colonnes += ", ..."
        n_lignes = len(df)
        n_colonnes = len(df.columns)
        lignes.append(
            f"- **{nom}** : {n_lignes} ligne{'s' if n_lignes != 1 else ''}, "
            f"{n_colonnes} colonne{'s' if n_colonnes != 1 else ''} ({apercu_colonnes})"
        )
    return "\n".join(lignes)


def detecter_tables_mentionnees(question: str, tables: dict) -> list[str]:
    """Renvoie les noms des tables chargees explicitement mentionnees dans la
    question (peu importe qu'elles viennent de fichiers separes ou de
    plusieurs feuilles d'un meme classeur : les deux sont traitees de facon
    identique une fois chargees dans `tables`).

    Reconnait aussi bien le nom exact charge qu'un alias informel (voir
    `alias_table`) qui omet le prefixe technique et/ou le singulier/pluriel
    - ex: "education" et "presence" reconnus pour "FNewEducation" et
    "FNewPresences", sans que l'equipe ait besoin de citer le nom technique
    complet.

    Pour les alias tres generiques qui se confondent avec le vocabulaire
    courant de l'observatoire (voir `ALIAS_AMBIGUS_GENERIQUES` - ex:
    "individus", "menages"), une mention nue ne suffit pas : il faut un
    ancrage explicite ("dans les individus", "table individus"...), sinon
    une question qui parle simplement "des individus" au sens general
    ferait croire, a tort, que la table opo_hypervel_individus est visee."""
    q = _sans_accents(question.lower())
    trouvees = []
    for nom in tables:
        if _sans_accents(nom.lower()) in q:
            trouvees.append(nom)
            continue
        aliases = alias_table(nom)
        aliases_specifiques = [a for a in aliases if a not in ALIAS_AMBIGUS_GENERIQUES]
        aliases_ambigus = [a for a in aliases if a in ALIAS_AMBIGUS_GENERIQUES]
        if any(_alias_mentionne(a, q) for a in aliases_specifiques):
            trouvees.append(nom)
        elif aliases_ambigus and any(_alias_ancre(a, q) for a in aliases_ambigus):
            trouvees.append(nom)
    return trouvees


def _colonnes_communes(df1: pd.DataFrame, df2: pd.DataFrame) -> list[str]:
    """Colonnes presentes dans les deux tables (comparaison insensible a la
    casse), dans l'ordre des colonnes de df1."""
    colonnes_b = {str(c).lower() for c in df2.columns}
    return [c for c in df1.columns if str(c).lower() in colonnes_b]


def _cle_str(serie: pd.Series) -> pd.Series:
    """Convertit une colonne de cle de jointure en texte, pour un merge entre
    deux tables ou le meme identifiant peut etre stocke avec un type
    different (ex: colonne entiere dans un export, texte dans un autre - cas
    reel rencontre : `individid` charge en int64 depuis une table, en texte
    depuis une autre, ce qui fait echouer `pd.merge` avec une ValueError
    explicite plutot qu'un simple resultat vide). Garde les valeurs
    manquantes manquantes (jamais la chaine litterale "nan", qui ferait sinon
    matcher a tort entre elles toutes les lignes sans identifiant renseigne)."""
    return serie.where(serie.isna(), serie.astype(str))


def _meilleure_cle_jointure(communes: list[str]) -> str | None:
    """Choisit, parmi des colonnes communes a deux tables, la seule fiable a
    utiliser comme cle de jointure - PAS simplement la premiere trouvee, et
    JAMAIS une colonne qui n'est pas documentee, meme en dernier recours.

    Bug reel corrige ici (episode 1) : `communes[0]` choisissait bien souvent
    "id" (quasi toujours la premiere colonne des vraies tables
    opo_hypervel_*), une cle primaire LOCALE a chaque table qui n'est jamais
    une reference vers une autre table dans ce schema.

    Bug reel corrige ici (episode 2, consigne explicite et repetee de
    l'observatoire - "id, menage_id, round_id, enquete_id... ne sont pas des
    identifiants a utiliser") : la version precedente de cette fonction
    retombait, a defaut d'identifiant confirme, sur "la premiere colonne
    commune qui n'est pas *id*" - ce qui choisissait encore silencieusement
    des colonnes Laravel/Hypervel du type `<table>_id` (`menage_id`,
    `round_id`, `enquete_id`, `individu_id`...) qui RESSEMBLENT a des
    identifiants par leur nom sans etre documentees comme de vraies cles.
    Ce filet de secours est supprime : sans identifiant confirme par le
    dictionnaire, il n'y a PAS de cle de jointure fiable, un point c'est
    tout - mieux vaut le dire clairement que deviner.

    Renvoie l'identifiant CONFIRME par le dictionnaire de donnees de
    l'observatoire (`IDENTIFIANTS_REELS_DOCUMENTES`) le plus tot dans l'ordre
    des colonnes communes, ou None si aucune des colonnes communes n'est
    confirmee (peu importe qu'elles s'appellent "id", "menage_id" ou
    n'importe quel autre nom qui ressemble a un identifiant)."""
    for c in communes:
        if str(c).strip().lower() in IDENTIFIANTS_REELS_DOCUMENTES:
            return c
    return None


def detecter_cles_communes(tables: dict) -> dict[tuple[str, str], list[str]]:
    """Detecte, pour chaque paire de tables chargees, les colonnes presentes
    dans les deux (candidates comme cle de jointure) - a partir des vraies
    donnees chargees, pas seulement de la documentation du dictionnaire."""
    noms = list(tables.keys())
    resultat = {}
    for i in range(len(noms)):
        for j in range(i + 1, len(noms)):
            a, b = noms[i], noms[j]
            communes = _colonnes_communes(tables[a], tables[b])
            if communes:
                resultat[(a, b)] = communes
    return resultat


def relation_entre_tables(nom1: str, nom2: str, tables: dict) -> str:
    """Decrit, a partir des vraies donnees chargees, comment deux tables sont
    reliees (colonnes en commun, candidates comme cle de jointure)."""
    manquantes = [n for n in (nom1, nom2) if n not in tables]
    if manquantes:
        return f"Table introuvable parmi celles chargées : {', '.join(manquantes)}."

    communes = _colonnes_communes(tables[nom1], tables[nom2])
    if not communes:
        return (
            f"Aucune colonne commune détectée entre **{nom1}** et **{nom2}** dans les données "
            "chargées : impossible de déterminer un lien direct entre ces deux tables telles quelles."
        )

    lignes = [
        f"**{nom1}** et **{nom2}** partagent {len(communes)} colonne(s) : "
        + ", ".join(f"`{c}`" for c in communes) + "."
    ]
    meilleure = _meilleure_cle_jointure(communes)
    if meilleure is not None:
        # `_meilleure_cle_jointure` ne renvoie plus JAMAIS qu'un identifiant
        # confirme par le dictionnaire (voir sa docstring) - cette colonne
        # l'est donc forcement.
        lignes.append(
            f"`{meilleure}` est la clé de jointure à utiliser (identifiant confirmé par le "
            f"dictionnaire de données). Demande « fusionne {nom1} et {nom2} » pour obtenir une "
            "table combinée."
        )
    else:
        non_confirmees = ", ".join(f"`{c}`" for c in communes)
        lignes.append(
            f"Aucune de ces colonnes communes n'est un identifiant confirmé par le dictionnaire de "
            f"données de l'observatoire ({non_confirmees}) : un nom de colonne qui ressemble à un "
            "identifiant (ex: `id`, `menage_id`, `round_id`, `enquete_id`...) n'en est pas la preuve, "
            "et ne doit jamais être utilisé comme clé de jointure sans confirmation. Vérifie auprès de "
            "l'équipe/du dictionnaire de données quel identifiant relie réellement ces deux tables."
        )
    return "\n".join(lignes)


def rapport_relations(tables: dict) -> str:
    """Resume les relations detectees (colonnes communes) entre toutes les
    paires de tables actuellement chargees (fichiers separes ou feuilles d'un
    meme classeur Excel : traites de la meme facon)."""
    if len(tables) < 2:
        return "Il faut au moins deux tables chargées pour détecter une relation entre elles."

    communes = detecter_cles_communes(tables)
    if not communes:
        return "Aucune colonne commune détectée entre les tables actuellement chargées."

    lignes = ["Relations détectées entre les tables chargées (colonnes en commun) :"]
    for (a, b), cols in communes.items():
        lignes.append(f"- **{a}** ↔ **{b}** : " + ", ".join(f"`{c}`" for c in cols))
    return "\n".join(lignes)


def fusionner_tables(nom1: str, nom2: str, tables: dict, cle: str | None = None) -> pd.DataFrame:
    """Fusionne (jointure) deux tables chargees sur une colonne commune.

    Si `cle` n'est pas precisee, la premiere colonne commune detectee entre
    les deux tables est utilisee automatiquement. Leve une ValueError si les
    tables sont introuvables, si aucune colonne commune n'existe, ou si la
    colonne demandee n'existe pas dans les deux tables."""
    manquantes = [n for n in (nom1, nom2) if n not in tables]
    if manquantes:
        raise ValueError(f"Table introuvable parmi celles chargées : {', '.join(manquantes)}")

    df1, df2 = tables[nom1], tables[nom2]

    if cle is None:
        communes = _colonnes_communes(df1, df2)
        if not communes:
            raise ValueError(f"Aucune colonne commune trouvée entre '{nom1}' et '{nom2}' pour fusionner.")
        cle = _meilleure_cle_jointure(communes)
        if cle is None:
            non_confirmees = ", ".join(f"'{c}'" for c in communes)
            raise ValueError(
                f"Aucune colonne commune CONFIRMÉE par le dictionnaire de données trouvée entre "
                f"'{nom1}' et '{nom2}' pour fusionner ({non_confirmees} sont communes, mais aucune "
                "n'est un identifiant documenté — un nom qui ressemble à un identifiant, comme 'id' "
                "ou 'menage_id', n'en est pas la preuve)."
            )
    elif cle not in df1.columns or cle not in df2.columns:
        raise ValueError(f"La colonne '{cle}' n'existe pas dans les deux tables.")

    return df1.merge(df2, on=cle, how="inner", suffixes=(f"_{nom1}", f"_{nom2}"))


def detecter_cle_jointure(nom1: str, nom2: str, tables: dict) -> str | None:
    """Renvoie la colonne commune utilisee (ou qui serait utilisee) comme cle
    de jointure entre deux tables chargees, ou None si aucune ne convient -
    utilise a la fois par le calcul reel (fusion/difference) et par la
    generation de syntaxe R/Stata, pour rester coherent sur la cle choisie."""
    if nom1 not in tables or nom2 not in tables:
        return None
    communes = _colonnes_communes(tables[nom1], tables[nom2])
    return _meilleure_cle_jointure(communes)


def difference_tables(nom1: str, nom2: str, tables: dict, cle: str | None = None) -> pd.DataFrame:
    """Lignes de la table `nom1` dont la cle n'a AUCUNE correspondance dans la
    table `nom2` (anti-jointure / difference d'ensembles) : repond a une
    question du type "qui est present dans X mais pas dans Y" (ex: individus
    enregistres sur une fiche mais absents d'une autre).

    Si `cle` n'est pas precisee, la premiere colonne commune detectee entre
    les deux tables est utilisee automatiquement (voir `detecter_cle_jointure`).
    Leve une ValueError si les tables sont introuvables, si aucune colonne
    commune n'existe, ou si la colonne demandee n'existe pas dans les deux
    tables."""
    manquantes = [n for n in (nom1, nom2) if n not in tables]
    if manquantes:
        raise ValueError(f"Table introuvable parmi celles chargées : {', '.join(manquantes)}")

    df1, df2 = tables[nom1], tables[nom2]

    if cle is None:
        cle = detecter_cle_jointure(nom1, nom2, tables)
        if cle is None:
            communes = _colonnes_communes(df1, df2)
            if not communes:
                raise ValueError(f"Aucune colonne commune trouvée entre '{nom1}' et '{nom2}' pour comparer.")
            non_confirmees = ", ".join(f"'{c}'" for c in communes)
            raise ValueError(
                f"Aucune colonne commune CONFIRMÉE par le dictionnaire de données trouvée entre "
                f"'{nom1}' et '{nom2}' pour comparer ({non_confirmees} sont communes, mais aucune "
                "n'est un identifiant documenté — un nom qui ressemble à un identifiant, comme 'id' "
                "ou 'menage_id', n'en est pas la preuve)."
            )
    elif cle not in df1.columns or cle not in df2.columns:
        raise ValueError(f"La colonne '{cle}' n'existe pas dans les deux tables.")

    fusion = df1.merge(df2[[cle]].drop_duplicates(), on=cle, how="left", indicator=True)
    return fusion[fusion["_merge"] == "left_only"].drop(columns="_merge")


def syntaxe_fusion(nom1: str, nom2: str, cle: str) -> str:
    """Syntaxe R et Stata equivalente a une fusion (jointure) entre deux
    tables sur une cle commune, a fournir en complement du resultat direct."""
    r = f'resultat <- merge({nom1}, {nom2}, by = "{cle}")'
    stata = f"use {nom1}, clear\nmerge 1:1 {cle} using {nom2}"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def syntaxe_difference(nom1: str, nom2: str, cle: str) -> str:
    """Syntaxe R (dplyr::anti_join) et Stata equivalente a une difference
    d'ensembles (lignes de nom1 sans correspondance dans nom2)."""
    r = f'resultat <- dplyr::anti_join({nom1}, {nom2}, by = "{cle}")'
    stata = f"use {nom1}, clear\nmerge 1:1 {cle} using {nom2}\nkeep if _merge == 1"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def tableau_croise(df: pd.DataFrame, colonne1: str, colonne2: str) -> pd.DataFrame:
    """Tableau croise (analyse bivariee) entre deux colonnes d'une meme table :
    effectifs croises, avec une ligne/colonne "Total" en marge - repond a une
    question du type "croise sex et education_level" ou "tableau croise entre
    X et Y"."""
    for c in (colonne1, colonne2):
        if c not in df.columns:
            raise ValueError(f"Colonne '{c}' absente. Colonnes disponibles : {list(df.columns)}")
    return pd.crosstab(df[colonne1], df[colonne2], margins=True, margins_name="Total")


def syntaxe_tableau_croise(nom_table: str, colonne1: str, colonne2: str) -> str:
    """Syntaxe R et Stata equivalente a un tableau croise entre deux colonnes."""
    r = f"table({nom_table}${colonne1}, {nom_table}${colonne2})"
    stata = f"use {nom_table}, clear\ntab {colonne1} {colonne2}"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def colonnes_numeriques(df: pd.DataFrame) -> list[str]:
    """Colonnes de type numerique (candidates a une correlation/analyse
    multivariee quantitative) - a partir du type reel des donnees chargees."""
    return [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]


def matrice_correlation(df: pd.DataFrame, colonnes: list[str] | None = None) -> pd.DataFrame:
    """Matrice de correlation (Pearson) entre colonnes numeriques - analyse
    multivariee simple pour reperer des liens entre plusieurs variables
    quantitatives a la fois. Si `colonnes` n'est pas precise, utilise toutes
    les colonnes numeriques detectees automatiquement."""
    cols = colonnes if colonnes is not None else colonnes_numeriques(df)
    cols_valides = [c for c in cols if c in df.columns and pd.api.types.is_numeric_dtype(df[c])]
    if len(cols_valides) < 2:
        raise ValueError(
            "Il faut au moins deux colonnes numériques pour calculer une corrélation "
            f"(colonnes numériques disponibles : {colonnes_numeriques(df)})."
        )
    return df[cols_valides].corr().round(3)


def syntaxe_correlation(nom_table: str, colonnes: list[str]) -> str:
    """Syntaxe R et Stata equivalente a une matrice de correlation."""
    vecteur_r = ", ".join(f'"{c}"' for c in colonnes)
    r = f"cor({nom_table}[, c({vecteur_r})], use = \"pairwise.complete.obs\")"
    stata = f"use {nom_table}, clear\npwcorr {' '.join(colonnes)}, sig"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def tableau_multivarie(df: pd.DataFrame, colonnes: list[str]) -> pd.DataFrame:
    """Effectifs croises sur 3 colonnes ou plus (analyse multivariee
    categorielle simple) - groupement par toutes les colonnes demandees,
    avec le nombre de lignes pour chaque combinaison observee."""
    manquantes = [c for c in colonnes if c not in df.columns]
    if manquantes:
        raise ValueError(f"Colonne(s) absente(s) : {', '.join(manquantes)}. Colonnes disponibles : {list(df.columns)}")
    return (
        df.groupby(colonnes, dropna=False)
        .size()
        .reset_index(name="effectif")
        .sort_values("effectif", ascending=False)
        .reset_index(drop=True)
    )


def syntaxe_tableau_multivarie(nom_table: str, colonnes: list[str]) -> str:
    """Syntaxe R et Stata equivalente a un groupement sur plusieurs colonnes."""
    vecteur_r = ", ".join(colonnes)
    r = f"resultat <- dplyr::count({nom_table}, {vecteur_r})"
    stata = f"use {nom_table}, clear\negen _grp = group({' '.join(colonnes)})\ntab _grp"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def syntaxe_repartition(nom_table: str, colonne: str) -> str:
    """Syntaxe R et Stata equivalente a une repartition (comptage/pourcentages)
    d'une colonne - a fournir en complement du resultat direct, pour que
    l'equipe puisse reproduire ou approfondir le calcul dans R/Stata."""
    r = f'table({nom_table}${colonne}); prop.table(table({nom_table}${colonne})) * 100'
    stata = f"use {nom_table}, clear\ntab {colonne}"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def syntaxe_echantillon(nom_table: str, n: int, seed: int) -> str:
    """Syntaxe R et Stata equivalente a un echantillon aleatoire reproductible
    (meme graine que celle utilisee par `echantillon`)."""
    r = f'set.seed({seed}); resultat <- {nom_table}[sample(nrow({nom_table}), {n}), ]'
    stata = f"use {nom_table}, clear\nset seed {seed}\nsample {n}, count"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def syntaxe_doublons(nom_table: str, colonne: str) -> str:
    """Syntaxe R et Stata equivalente a une detection de doublons sur une
    colonne d'identifiant."""
    r = (
        f"resultat <- {nom_table}[{nom_table}${colonne} %in% "
        f"{nom_table}${colonne}[duplicated({nom_table}${colonne})], ]"
    )
    stata = f"use {nom_table}, clear\nduplicates tag {colonne}, generate(_dup)\nlist if _dup > 0"
    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


def syntaxe_coherence(nom_table: str, colonnes_id: list[str], colonnes_date: list[str]) -> str:
    """Syntaxe R et Stata equivalente au rapport de coherence (doublons sur
    les colonnes d'ID detectees + dates invraisemblables sur les colonnes de
    date detectees) - combine les deux controles dans un seul bloc de code."""
    lignes_r = [f"use_table <- {nom_table}"]
    for c in colonnes_id:
        lignes_r.append(f'sum(duplicated(use_table${c}))  # doublons sur {c}')
    for c in colonnes_date:
        lignes_r.append(
            f'sum(is.na(as.Date(use_table${c})) | as.Date(use_table${c}) > Sys.Date())  # dates invraisemblables sur {c}'
        )
    r = "\n".join(lignes_r) if (colonnes_id or colonnes_date) else "# Aucune colonne d'identifiant ou de date detectee"

    lignes_stata = [f"use {nom_table}, clear"]
    for c in colonnes_id:
        lignes_stata.append(f"duplicates report {c}")
    for c in colonnes_date:
        lignes_stata.append(f"count if missing({c}) | {c} > date(\"$S_DATE\", \"DMY\")")
    stata = "\n".join(lignes_stata) if (colonnes_id or colonnes_date) else "# Aucune colonne d'identifiant ou de date detectee"

    return f"**Syntaxe équivalente :**\n```r\n{r}\n```\n```stata\n{stata}\n```"


"""
Catalogue de controles de coherence avances, specifiques au type d'enquete de
l'observatoire (au-dela des doublons/dates generiques de `rapport_coherence`).

Chaque controle est auto-detecte par le NOM des colonnes / tables reellement
chargees (memes principe que `detect_id_columns`/`detect_agent_columns`),
puisque les noms exacts varient d'un fichier a l'autre et ne sont jamais
codes en dur. Un controle dont les colonnes necessaires ne sont pas
detectees renvoie None (ignore proprement) plutot que de risquer un faux
positif sur des colonnes qui ne correspondent pas a ce qu'il est cense
verifier - transparence : le rapport final liste toujours explicitement quels
controles ont pu s'appliquer et lesquels ont ete ignores faute de colonnes
reconnues.
"""

BIRTH_DATE_LIKE = re.compile(r"(birth_date|naissance)", re.IGNORECASE)
DEATH_DATE_LIKE = re.compile(r"(death_date|evtdate|deces|dcd)", re.IGNORECASE)
ARRIVE_DATE_LIKE = re.compile(r"arrive.*date|date.*arrive", re.IGNORECASE)
DEPART_DATE_LIKE = re.compile(r"depart.*date|date.*depart", re.IGNORECASE)
ENTRY_DATE_LIKE = re.compile(r"(entry_date|date.*enregistr)", re.IGNORECASE)
WEIGHT_LIKE = re.compile(r"(poids|weight)", re.IGNORECASE)
HEIGHT_LIKE = re.compile(r"(taille|height)", re.IGNORECASE)
MOTHERID_LIKE = re.compile(r"(motherid|mereid|id_mere)", re.IGNORECASE)
FATHERID_LIKE = re.compile(r"(fatherid|pereid|id_pere)", re.IGNORECASE)
LAT_LIKE = re.compile(r"^lat(itude)?$", re.IGNORECASE)
LON_LIKE = re.compile(r"^lon(g|gitude)?$", re.IGNORECASE)
PHONE_LIKE = re.compile(r"(telephone|numtel|phone)", re.IGNORECASE)
SLEEP_LIKE = re.compile(r"(sleep|dormi)", re.IGNORECASE)
LOCATION_LIKE = re.compile(r"(locationid|menageid)", re.IGNORECASE)
SOCIALGP_LIKE = re.compile(r"(socialgpid)", re.IGNORECASE)
UNION_AGE_LIKE = re.compile(r"union_?age", re.IGNORECASE)
DIEDINHS_LIKE = re.compile(r"diedinhs", re.IGNORECASE)
GONETOHS_LIKE = re.compile(r"gonetohs", re.IGNORECASE)
RLTN_HEAD_LIKE = re.compile(r"rltn_?head", re.IGNORECASE)
# "1 Socialgp head" (chef de menage) : code documente tel quel dans le
# dictionnaire de donnees de l'observatoire ("source_Dictionnaire des
# variables.txt", colonne `rltn_head`) - jamais devine.
CODE_CHEF_MENAGE = "1"
PREGOUTID_LIKE = re.compile(r"pregoutid", re.IGNORECASE)
EVENTID_LIKE = re.compile(r"^eventid$", re.IGNORECASE)
PEVENTID_LIKE = re.compile(r"^peventid$", re.IGNORECASE)
CPN_DATE_LIKE = re.compile(r"^cpn_?date(\d+)$", re.IGNORECASE)
NB_CPN_LIKE = re.compile(r"nb_?cpn", re.IGNORECASE)
RESPONDID_LIKE = re.compile(r"^respondid$", re.IGNORECASE)
S4_2_LIKE = re.compile(r"^S4_2$", re.IGNORECASE)
S4_2MM_LIKE = re.compile(r"^S4_2mm$", re.IGNORECASE)
S5_1_LIKE = re.compile(r"^S5_1$", re.IGNORECASE)
S5_2_LIEU_LIKE = re.compile(r"^S5_2[ABC]$", re.IGNORECASE)
RES_STATUS_LIKE = re.compile(r"res_?status", re.IGNORECASE)
BEGIN_TIME_LIKE = re.compile(r"begin_?time", re.IGNORECASE)
END_TIME_LIKE = re.compile(r"end_?time", re.IGNORECASE)
# "1. Resident" : code documente tel quel dans le dictionnaire de donnees de
# l'observatoire ("source_Dictionnaire des variables.txt", colonne
# `res_status`) - jamais devine.
CODE_RESIDENT = "1"
# "1. Oui" / "2. Non" : codes documentes tels quels dans le questionnaire
# source "20_FICHE SANTE_R14" (question S4_2) - jamais devines.
CODE_NON = "2"
GENDER_LIKE = re.compile(r"^gender$", re.IGNORECASE)
# "1. Male" / "2. Female" : codes documentes tels quels dans le dictionnaire
# de donnees de l'observatoire ("source_Dictionnaire des variables.txt",
# colonne `gender`) - jamais devines.
CODE_HOMME = "1"
LIVING_CHILDREN_LIKE = re.compile(r"living_?children_?number", re.IGNORECASE)
ISALIVE_LIKE = re.compile(r"^isAlive$", re.IGNORECASE)
BIRTHDATE_ENFANT_LIKE = re.compile(r"^birthDate$", re.IGNORECASE)
# Les 4 variantes de date d'union documentees sur FNewRelationship (schema
# reel de l'observatoire) : debut, civile, religieuse, traditionnelle -
# chacune doit rester posterieure a la naissance de l'individu.
DATES_UNION_LIKE = {
    "début d'union": re.compile(r"uni_?start_?date", re.IGNORECASE),
    "union civile": re.compile(r"uni_?civil_?date", re.IGNORECASE),
    "union religieuse": re.compile(r"uni_?relig_?date", re.IGNORECASE),
    "union traditionnelle": re.compile(r"uni_?trad_?date", re.IGNORECASE),
}

# Bornes approximatives du territoire burkinabe, pour reperer des coordonnees
# GPS clairement hors zone (au-dela d'une simple valeur manquante).
BF_LAT_MIN, BF_LAT_MAX = 9.0, 15.2
BF_LON_MIN, BF_LON_MAX = -5.6, 2.5


def _premiere_colonne(df: pd.DataFrame, motif: re.Pattern) -> str | None:
    for c in df.columns:
        if motif.search(str(c)):
            return c
    return None


def _age_en_annees(df: pd.DataFrame, colonne_naissance: str) -> pd.Series:
    naissance = pd.to_datetime(df[colonne_naissance], errors="coerce", dayfirst=True)
    return (pd.Timestamp(datetime.now()) - naissance).dt.days / 365.25


def controle_id_longueur(df: pd.DataFrame) -> dict | None:
    """Signale les identifiants dont la longueur (nombre de caracteres)
    differe de la longueur la plus frequente dans la meme colonne - repere un
    ID mal saisi sans connaitre a l'avance la longueur exacte attendue."""
    colonnes_id = detect_id_columns(df)
    if not colonnes_id:
        return None
    total = 0
    exemples = []
    for col in colonnes_id:
        non_nuls = df[col].dropna()
        if non_nuls.empty:
            continue
        longueur_attendue = int(non_nuls.astype(str).str.len().mode().iloc[0])
        longueurs_completes = df[col].astype(str).str.len()
        mauvaises_mask = df[col].notna() & (longueurs_completes != longueur_attendue)
        n = int(mauvaises_mask.sum())
        if n > 0:
            total += n
            exemples.append(f"{col} (longueur attendue {longueur_attendue})")
    return {"colonnes_verifiees": colonnes_id, "n_anomalies": total, "detail": exemples}


def controle_auto_reference(df: pd.DataFrame) -> dict | None:
    """Signale les lignes ou un identifiant se reference lui-meme (ex:
    individid == individid2), a partir de toute paire de colonnes ID detectee
    dont l'une se termine par "2" par rapport a l'autre."""
    colonnes_id = set(detect_id_columns(df))
    verifiees, total = [], 0
    for col in colonnes_id:
        if col.endswith("2"):
            continue
        col2 = col + "2"
        if col2 not in df.columns:
            continue
        verifiees.append((col, col2))
        n = int((df[col] == df[col2]).sum())
        total += n
    if not verifiees:
        return None
    return {"colonnes_verifiees": [f"{a}/{b}" for a, b in verifiees], "n_anomalies": total}


def controle_parents_identiques(df: pd.DataFrame) -> dict | None:
    """Signale les lignes ou fatherid == motherid (parents identiques)."""
    col_pere = _premiere_colonne(df, FATHERID_LIKE)
    col_mere = _premiere_colonne(df, MOTHERID_LIKE)
    if col_pere is None or col_mere is None:
        return None
    masque = df[col_pere].notna() & (df[col_pere] == df[col_mere])
    return {"colonnes_verifiees": [col_pere, col_mere], "n_anomalies": int(masque.sum())}


def controle_parent_manquant_jeune_enfant(df: pd.DataFrame, seuil_age: float = 5.0) -> dict | None:
    """Signale les moins de `seuil_age` ans sans motherid renseigne, et
    separement sans motherid NI fatherid renseignes."""
    col_naissance = _premiere_colonne(df, BIRTH_DATE_LIKE)
    col_mere = _premiere_colonne(df, MOTHERID_LIKE)
    if col_naissance is None or col_mere is None:
        return None
    age = _age_en_annees(df, col_naissance)
    jeunes = age < seuil_age
    sans_mere = int((jeunes & df[col_mere].isna()).sum())
    resultat = {
        "colonnes_verifiees": [col_naissance, col_mere],
        "n_anomalies": sans_mere,
        "detail": [f"< {seuil_age} ans sans motherid : {sans_mere}"],
    }
    col_pere = _premiere_colonne(df, FATHERID_LIKE)
    if col_pere is not None:
        # Sous-ensemble de `sans_mere` (deja compte dans n_anomalies) : indication
        # supplementaire, ne pas re-additionner pour eviter un double comptage.
        sans_les_deux = int((jeunes & df[col_mere].isna() & df[col_pere].isna()).sum())
        resultat["colonnes_verifiees"].append(col_pere)
        resultat["detail"].append(f"< {seuil_age} ans sans motherid ni fatherid : {sans_les_deux}")
    return resultat


def controle_dates_ordre(df: pd.DataFrame, motif_avant: re.Pattern, motif_apres: re.Pattern, strict: bool = True) -> dict | None:
    """Controle generique : signale les lignes ou une date qui devrait
    logiquement precéder une autre lui est posterieure (ex: naissance apres
    deces, enregistrement avant naissance). `motif_avant` doit designer la
    date qui devrait etre la plus ancienne."""
    col_avant = _premiere_colonne(df, motif_avant)
    col_apres = _premiere_colonne(df, motif_apres)
    if col_avant is None or col_apres is None or col_avant == col_apres:
        return None
    d_avant = pd.to_datetime(df[col_avant], errors="coerce", dayfirst=True)
    d_apres = pd.to_datetime(df[col_apres], errors="coerce", dayfirst=True)
    masque = (d_avant > d_apres) if strict else (d_avant >= d_apres)
    masque = masque.fillna(False)
    return {"colonnes_verifiees": [col_avant, col_apres], "n_anomalies": int(masque.sum())}


def controle_sentinelle(df: pd.DataFrame, motif_colonne: re.Pattern, valeur_sentinelle) -> dict | None:
    """Signale les lignes ou une colonne (ex: poids, taille) contient une
    valeur sentinelle connue (9999, 99...) indiquant une non-reponse codee
    plutot qu'une vraie mesure."""
    col = _premiere_colonne(df, motif_colonne)
    if col is None:
        return None
    masque = pd.to_numeric(df[col], errors="coerce") == valeur_sentinelle
    return {"colonnes_verifiees": [col], "n_anomalies": int(masque.sum())}


def controle_gps_hors_zone(df: pd.DataFrame) -> dict | None:
    """Signale les coordonnees GPS manquantes ou hors du territoire
    burkinabe (bornes approximatives)."""
    col_lat = _premiere_colonne(df, LAT_LIKE)
    col_lon = _premiere_colonne(df, LON_LIKE)
    if col_lat is None or col_lon is None:
        return None
    lat = pd.to_numeric(df[col_lat], errors="coerce")
    lon = pd.to_numeric(df[col_lon], errors="coerce")
    hors_zone = (
        lat.isna() | lon.isna()
        | (lat < BF_LAT_MIN) | (lat > BF_LAT_MAX)
        | (lon < BF_LON_MIN) | (lon > BF_LON_MAX)
    )
    return {"colonnes_verifiees": [col_lat, col_lon], "n_anomalies": int(hors_zone.sum())}


def controle_telephone_format(df: pd.DataFrame) -> dict | None:
    """Signale les numeros de telephone dont un des segments (separes par
    "/") ne fait pas 8 chiffres (format attendu au Burkina Faso)."""
    col = _premiere_colonne(df, PHONE_LIKE)
    if col is None:
        return None

    def _invalide(valeur) -> bool:
        if pd.isna(valeur):
            return False
        segments = str(valeur).split("/")
        return any(not seg.strip().isdigit() or len(seg.strip()) != 8 for seg in segments if seg.strip())

    masque = df[col].apply(_invalide)
    return {"colonnes_verifiees": [col], "n_anomalies": int(masque.sum())}


def controle_dates_arrivee_depart(df: pd.DataFrame) -> dict | None:
    """Fiche presence : signale les dates de depart anterieures ou egales a
    la date d'arrivee (incoherentes ou identiques)."""
    col_arrivee = _premiere_colonne(df, ARRIVE_DATE_LIKE)
    col_depart = _premiere_colonne(df, DEPART_DATE_LIKE)
    if col_arrivee is None or col_depart is None:
        return None
    arrivee = pd.to_datetime(df[col_arrivee], errors="coerce", dayfirst=True)
    depart = pd.to_datetime(df[col_depart], errors="coerce", dayfirst=True)
    masque = ((depart < arrivee) | (depart == arrivee)).fillna(False)
    return {"colonnes_verifiees": [col_arrivee, col_depart], "n_anomalies": int(masque.sum())}


def controle_coherence_presence(df: pd.DataFrame) -> list[tuple[str, dict]] | None:
    """Fiche presence (FNewPresences) : 3 controles de coherence entre le
    statut "a dormi sur place" et les dates d'arrivee/depart renseignees -
    - a dormi (sleep=oui) MAIS une date de depart est quand meme renseignee
      (incoherent : si la personne a dormi sur place, elle n'est pas censee
      etre partie) ;
    - n'a PAS dormi (sleep=non) SANS aucune date de depart renseignee
      (incoherent : une absence devrait normalement etre expliquee par un
      depart date) ;
    - ni date de depart ni date d'arrivee renseignees du tout (fiche
      incomplete sur le statut de presence, distinct du cas "dates identiques"
      deja couvert par `controle_dates_arrivee_depart`).

    Renvoie une liste de (libelle, resultat), ou None si les colonnes
    necessaires (sleep + au moins une des deux dates) ne sont pas
    detectees."""
    col_sleep = _premiere_colonne(df, SLEEP_LIKE)
    col_arrivee = _premiere_colonne(df, ARRIVE_DATE_LIKE)
    col_depart = _premiere_colonne(df, DEPART_DATE_LIKE)
    if col_sleep is None or (col_arrivee is None and col_depart is None):
        return None

    sleep_norm = df[col_sleep].astype(str).str.strip().str.lower()
    sleep_oui = sleep_norm.isin(["1", "1.0", "oui", "yes", "true"])
    sleep_non = df[col_sleep].notna() & ~sleep_oui

    resultats: list[tuple[str, dict]] = []
    if col_depart is not None:
        a_depart = df[col_depart].notna()
        resultats.append((
            "A dormi sur place mais une date de départ est renseignée",
            {"colonnes_verifiees": [col_sleep, col_depart], "n_anomalies": int((sleep_oui & a_depart).sum())},
        ))
        resultats.append((
            "N'a pas dormi sur place sans date de départ renseignée",
            {"colonnes_verifiees": [col_sleep, col_depart], "n_anomalies": int((sleep_non & ~a_depart).sum())},
        ))
    if col_arrivee is not None and col_depart is not None:
        masque_incomplete = df[col_arrivee].isna() & df[col_depart].isna()
        resultats.append((
            "Ni date de départ ni date d'arrivée renseignées",
            {"colonnes_verifiees": [col_arrivee, col_depart], "n_anomalies": int(masque_incomplete.sum())},
        ))
    return resultats or None


def controle_residence_multiple(df: pd.DataFrame) -> dict | None:
    """Signale un individu present dans plusieurs menages/localisations
    differents (colonne `locationid` ou `socialgpid` associee a plusieurs
    valeurs distinctes pour le meme individid)."""
    colonnes_id = detect_id_columns(df)
    col_individu = next((c for c in colonnes_id if "individ" in c.lower()), None)
    col_lieu = _premiere_colonne(df, LOCATION_LIKE) or _premiere_colonne(df, SOCIALGP_LIKE)
    if col_individu is None or col_lieu is None or col_individu == col_lieu:
        return None
    nb_lieux_distincts = df.groupby(col_individu)[col_lieu].nunique()
    individus_multiples = nb_lieux_distincts[nb_lieux_distincts > 1]
    return {"colonnes_verifiees": [col_individu, col_lieu], "n_anomalies": int(len(individus_multiples))}


def controle_tranche_age(df: pd.DataFrame, borne_min: float, borne_max: float) -> dict | None:
    """Controle generique : signale les lignes dont l'age (calcule depuis la
    colonne de naissance detectee) sort d'une tranche attendue pour ce type
    de fiche (ex: 12-49 ans pour une fiche genesique, 5-34 ans pour une fiche
    education...). Ne fonctionne que si LA TABLE ELLE-MEME porte une colonne
    de naissance - voir `controle_tranche_age_croisee` pour le cas (largement
    majoritaire sur le vrai schema de l'observatoire) ou ce n'est pas le cas."""
    col = _premiere_colonne(df, BIRTH_DATE_LIKE)
    if col is None:
        return None
    age = _age_en_annees(df, col)
    hors_plage = ((age < borne_min) | (age > borne_max)).fillna(False)
    return {
        "colonnes_verifiees": [col],
        "n_anomalies": int(hors_plage.sum()),
        "detail": [f"âge hors {borne_min:g}-{borne_max:g} ans"],
    }


def _table_individus(tables: dict, exclure: str | None = None) -> str | None:
    """Trouve, parmi les tables chargees, celle qui porte l'identite de
    reference de chaque individu - `individid` + une colonne de naissance
    (FNewIndividual/RegAllIndividual dans le schema reel de l'observatoire).

    Necessaire car, contrairement a ce qu'on pourrait supposer, PRESQUE
    AUCUNE fiche d'evenement du vrai schema (education, emploi, grossesse,
    genesique, histoire matrimoniale, telephone...) ne porte elle-meme de
    colonne birth_date : elle vit UNIQUEMENT sur la table individus, reliee
    par `individid` - verifie en inspectant les 27 vraies tables exportees le
    2026-08-17. Sans ce croisement, `controle_tranche_age` (qui ne cherche
    birth_date QUE dans la table en cours) ne peut jamais s'appliquer a
    aucune de ces fiches, et le controle d'age correspondant est ignore en
    silence sur les vraies donnees (jamais signale comme un vrai bug avant
    inspection directe du schema)."""
    return next(
        (
            n for n, df in tables.items()
            if n != exclure
            and _table_correspond(n, ["individual", "individu"])
            and any("individ" in c.lower() and not c.lower().endswith("2") for c in detect_id_columns(df))
            and _premiere_colonne(df, BIRTH_DATE_LIKE) is not None
        ),
        None,
    )


def controle_tranche_age_croisee(
    df: pd.DataFrame, tables: dict, nom_table: str, borne_min: float, borne_max: float
) -> dict | None:
    """Version croisee de `controle_tranche_age`, a utiliser en priorite pour
    les controles d'age par type de fiche (`TRANCHES_AGE_PAR_TYPE_FICHE`) :
    tente d'abord la colonne de naissance LOCALE (cas rare), puis, si absente,
    va chercher l'age via une jointure `individid` -> la table individus
    (voir `_table_individus`) plutot que d'abandonner. Repond a l'age ACTUEL
    de la personne (date du jour - naissance), pertinent pour les controles
    d'eligibilite/tranche d'age courants du catalogue (education, emploi,
    genesique, grossesse...) - PAS a un age a une date d'evenement passee
    (ex: age a la premiere union), qui necessite un calcul dedie (voir
    `controle_age_union`)."""
    resultat_local = controle_tranche_age(df, borne_min, borne_max)
    if resultat_local is not None:
        return resultat_local

    col_id = next((c for c in detect_id_columns(df) if "individ" in c.lower() and not c.lower().endswith("2")), None)
    if col_id is None or col_id not in df.columns:
        return None
    nom_individus = _table_individus(tables, exclure=nom_table)
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_naissance = _premiere_colonne(individus, BIRTH_DATE_LIKE)
    if col_id_individus is None or col_naissance is None:
        return None

    reference = individus[[col_id_individus, col_naissance]].rename(columns={col_id_individus: col_id})
    reference = reference.assign(**{col_id: _cle_str(reference[col_id])})
    gauche = df[[col_id]].assign(**{col_id: _cle_str(df[col_id])})
    fusion = gauche.merge(reference, on=col_id, how="left")
    age = _age_en_annees(fusion, col_naissance)
    hors_plage = ((age < borne_min) | (age > borne_max)).fillna(False)
    return {
        "colonnes_verifiees": [col_id, f"{nom_individus}.{col_naissance}"],
        "n_anomalies": int(hors_plage.sum()),
        "detail": [f"âge (via {col_id} → {nom_individus}) hors {borne_min:g}-{borne_max:g} ans"],
    }


CONTROLES_GENERIQUES_PAR_TABLE = [
    ("Identifiants de longueur inhabituelle", controle_id_longueur),
    ("Auto-référence (ID == ID2)", controle_auto_reference),
    ("Parents identiques (fatherid == motherid)", controle_parents_identiques),
    ("Jeune enfant sans parent renseigné", controle_parent_manquant_jeune_enfant),
    ("Valeur de poids sentinelle (9999)", lambda df: controle_sentinelle(df, WEIGHT_LIKE, 9999)),
    ("Valeur de taille sentinelle (99)", lambda df: controle_sentinelle(df, HEIGHT_LIKE, 99)),
    ("Coordonnées GPS manquantes ou hors zone", controle_gps_hors_zone),
    ("Format de téléphone invalide", controle_telephone_format),
    ("Dates arrivée/départ incohérentes", controle_dates_arrivee_depart),
    ("Résidence multiple pour un même individu", controle_residence_multiple),
    ("Naissance postérieure au décès", lambda df: controle_dates_ordre(df, BIRTH_DATE_LIKE, DEATH_DATE_LIKE)),
    # Bug reel corrige ici (jamais teste jusqu'ici) : le libelle annonce
    # detecter un ENREGISTREMENT anterieur a la naissance (entry_date <
    # birth_date, logiquement impossible - on ne peut pas enregistrer une
    # naissance avant qu'elle n'ait eu lieu), mais l'appel precedent avait
    # motif_avant/motif_apres INVERSES (ENTRY_DATE_LIKE, BIRTH_DATE_LIKE) :
    # `controle_dates_ordre` flagge quand `motif_avant` (cense etre le plus
    # ancien) se retrouve APRES `motif_apres`, donc l'ancien appel flaggait
    # en realite "entry_date >= birth_date" (le cas NORMAL, la quasi-totalite
    # des lignes), et laissait passer silencieusement le vrai cas anormal
    # (entry_date < birth_date). BIRTH_DATE_LIKE doit etre `motif_avant`
    # (la naissance doit toujours précéder l'enregistrement).
    ("Enregistrement antérieur à la naissance", lambda df: controle_dates_ordre(df, BIRTH_DATE_LIKE, ENTRY_DATE_LIKE)),
]

# Tranches d'age attendues par type de fiche (mots-cles reconnus dans le nom
# de la table charge - voir alias_table pour le meme principe de
# reconnaissance informelle des noms de table).
TRANCHES_AGE_PAR_TYPE_FICHE = [
    # "genesiqcompl"/"gnesiqcompl" : anciens noms synthetiques de test ;
    # "genesiquecomplementaire" : vrai nom reel (table
    # "opo_hypervel_histoire_genesique_complementaires", sans underscores
    # "histoiregenesiquecomplementaires") - garde les deux formes pour ne
    # jamais regresser sur les donnees de test existantes.
    (["genesiqcompl", "gnesiqcompl", "genesiquecomplementaire"], None),  # cas particulier traite a part (age de la mere)
    (["genesiq", "gnesiq"], (12.0, 49.0)),
    # "histmat"/"matrimon"/"union"/"marietal" : PAS d'entree ici - l'age
    # attendu (12-40 ans) concerne l'age A LA PREMIERE UNION, pas l'age
    # ACTUEL de la personne (que calculerait ce controle generique) : deux
    # valeurs tres differentes pour quelqu'un marie il y a des annees. Voir
    # `controle_age_union`, qui utilise directement la colonne `union_age`
    # declaree (FNewBase_HistMat) plutot que de deviner un age via naissance.
    (["education"], (5.0, 34.0)),
    (["emploi", "employ"], (15.0, 120.0)),
    (["pregnancy", "grossesse"], (12.0, 49.0)),
]


def _normaliser_nom_table(nom: str) -> str:
    """Version normalisee d'un nom de table pour la reconnaissance par
    mots-cles : prefixe technique retire (voir `PREFIXES_TABLE_IGNORES`) et
    UNDERSCORES SUPPRIMES. Ce deuxieme point est necessaire pour le schema
    reel de l'observatoire (base Hypervel), ou certains noms de table
    generes automatiquement epellent un sigle avec un underscore entre
    chaque lettre (ex: la table des deces s'appelle "opo_hypervel_d_e_c_e_s",
    celle des CPN "opo_hypervel_c_p_n_s") : sans ce retrait, aucun mot-cle
    ("deces", "cpn"...) ne peut jamais s'y retrouver comme sous-chaine
    contigue."""
    n = nom.lower()
    for prefixe in PREFIXES_TABLE_IGNORES:
        if n.startswith(prefixe) and len(n) > len(prefixe):
            n = n[len(prefixe):]
            break
    return n.replace("_", "")


def _table_correspond(nom: str, mots_cles: list[str]) -> bool:
    n = nom.lower()
    n_normalise = _normaliser_nom_table(nom)
    return any(m in n or m in n_normalise for m in mots_cles)


def trouver_table_par_role(tables: dict, mots_cles: list[str]) -> str | None:
    """Trouve la premiere table chargee dont le nom correspond a un role
    donne (ex: `["death", "deces"]` -> la table des deces, quel que soit son
    nom technique reel) - helper public (contrairement a `_table_correspond`)
    pour que les modules appelants (ex: app.py) puissent identifier une table
    par son role plutot que d'avoir a connaitre son nom exact."""
    return next((n for n in tables if _table_correspond(n, mots_cles)), None)


def rapport_coherence_avancee(tables: dict, nom_table: str | None = None) -> dict:
    """Execute le catalogue de controles de coherence avances sur une table
    precise (si `nom_table` est fourni) ou sur TOUTES les tables chargees,
    plus les controles croises entre tables (eligibilite presence <->
    education/emploi/genesique/pauvrete/sante, deces <-> presence, migration
    <-> presence, grossesse <-> issue de grossesse, snakebite menages <->
    residents).

    Chaque controle non applicable (colonnes non detectees) est explicitement
    liste comme "ignoré" plutot que d'etre tu, pour rester precis sur ce qui a
    reellement ete verifie - jamais de faux positif silencieux sur des
    colonnes qui ne correspondent pas a ce qui est cense etre controle."""
    tables_a_verifier = {nom_table: tables[nom_table]} if nom_table and nom_table in tables else tables

    resultats_par_table = {}
    for nom, df in tables_a_verifier.items():
        controles_ok, controles_ignores = [], []
        for libelle, fonction in CONTROLES_GENERIQUES_PAR_TABLE:
            resultat = fonction(df)
            if resultat is None:
                controles_ignores.append(libelle)
            else:
                controles_ok.append((libelle, resultat))

        for mots_cles, bornes in TRANCHES_AGE_PAR_TYPE_FICHE:
            if bornes is not None and _table_correspond(nom, mots_cles):
                # `tables` (l'ensemble COMPLET, pas `tables_a_verifier` qui
                # peut etre restreint a une seule table si `nom_table` est
                # precise) - le croisement vers la table individus a besoin
                # de voir toutes les tables chargees, meme quand on ne
                # verifie qu'une seule table cible.
                resultat = controle_tranche_age_croisee(df, tables, nom, *bornes)
                libelle = f"Âge hors tranche attendue ({bornes[0]:g}-{bornes[1]:g} ans)"
                if resultat is None:
                    controles_ignores.append(libelle)
                else:
                    controles_ok.append((libelle, resultat))

        if _table_correspond(nom, ["histmat", "matrimon", "union", "marietal"]):
            libelle = "Âge à la première union hors tranche attendue (12-40 ans)"
            resultat = controle_age_union(df)
            if resultat is None:
                controles_ignores.append(libelle)
            else:
                controles_ok.append((libelle, resultat))

        if _table_correspond(nom, ["presences", "presence"]):
            resultats_presence = controle_coherence_presence(df)
            if resultats_presence is None:
                controles_ignores.append("Cohérence dormi/dates d'arrivée-départ")
            else:
                controles_ok.extend(resultats_presence)

        if _table_correspond(nom, ["death", "deces"]):
            for libelle, resultat in (
                ("Individu déclaré décédé plus d'une fois", controle_doublon_individu(df)),
                ("Décédé en formation sanitaire sans y être jamais allé", controle_deces_formation_sanitaire(df)),
            ):
                if resultat is None:
                    controles_ignores.append(libelle)
                else:
                    controles_ok.append((libelle, resultat))

        if _table_correspond(nom, ["birth", "naissance"]):
            libelle = "Moins de 12 ans déclaré chef de ménage"
            resultat = controle_chef_menage_mineur(df)
            if resultat is None:
                controles_ignores.append(libelle)
            else:
                controles_ok.append((libelle, resultat))

        if _table_correspond(nom, ["cpn"]):
            for libelle, resultat in (
                ("Dates CPN désordonnées", controle_cpn_dates_desordonnees(df)),
                ("Dates CPN manquantes (nb_cpn déclaré > dates renseignées)", controle_cpn_dates_manquantes(df)),
            ):
                if resultat is None:
                    controles_ignores.append(libelle)
                else:
                    controles_ok.append((libelle, resultat))

        if _table_correspond(nom, ["gnesiqcompl", "genesiqcompl", "genesiquecomplementaire"]):
            libelle = "Année de naissance de l'enfant hors 2022-2026"
            resultat = controle_annee_naissance_enfant(df)
            if resultat is None:
                controles_ignores.append(libelle)
            else:
                controles_ok.append((libelle, resultat))

        if _table_correspond(nom, ["observation"]):
            libelle = "Durée d'entretien anormalement courte ou incohérente"
            resultat = controle_duree_entretien(df)
            if resultat is None:
                controles_ignores.append(libelle)
            else:
                controles_ok.append((libelle, resultat))

        if _table_correspond(nom, ["sante"]):
            for libelle, resultat in (
                ("Même répondant pour des ménages différents", controle_sante_doublon_menage(df)),
                ("Mois de dernière règle inconnu", controle_sante_mois_regle_inconnu(df)),
                ("Jamais utilisé internet mais utilisé (maison/travail/espace public)",
                 controle_sante_internet_contradictoire(df)),
            ):
                if resultat is None:
                    controles_ignores.append(libelle)
                else:
                    controles_ok.append((libelle, resultat))

        if controles_ok or controles_ignores:
            resultats_par_table[nom] = {"controles_ok": controles_ok, "controles_ignores": controles_ignores}

    # Controles croises entre tables (population eligible presence <->
    # tables associees), a partir d'une reconnaissance du role de chaque
    # table par mots-cles dans son nom.
    controles_croises = []
    nom_presence = next((n for n in tables if _table_correspond(n, ["presences", "presence"])), None)
    if nom_presence is not None:
        # Controles d'eligibilite PAR INDIVIDU (cle individid, commune a la
        # fiche presence et a la table cible).
        cibles = {
            "éducation": ["education"],
            "emploi": ["emploi", "employ"],
            "histoire génésique complémentaire": ["gnesiqcompl", "genesiqcompl", "genesiquecomplementaire"],
            "téléphone": ["telephone"],
        }
        # NOTE : "santé" utilise `respondid`, pas `individid` directement -
        # `controle_eligibilite_croisee` (base sur `detecter_cle_jointure`,
        # colonnes communes) ne le trouverait pas automatiquement comme cle
        # partagee avec la fiche presence. Le controle par menage
        # (cibles_par_menage, ci-dessous) couvre deja l'eligibilite pour
        # cette table ; `controle_sante_repondant_non_dormi` (branche plus
        # bas) couvre en plus le cas individuel via `respondid`.
        for libelle, mots in cibles.items():
            nom_cible = next((n for n in tables if _table_correspond(n, mots)), None)
            if nom_cible is None:
                continue
            resultat = controle_eligibilite_croisee(tables, nom_presence, nom_cible)
            if resultat is not None:
                controles_croises.append((f"Éligibilité présence ↔ {libelle}", nom_presence, nom_cible, resultat))

        # Controles d'eligibilite PAR MENAGE (cle socialgpid, absente de la
        # fiche presence : derivee via la table individus - voir
        # `controle_eligibilite_croisee_par_menage`).
        cibles_par_menage = {"pauvreté": ["pauvrete"], "santé": ["sante"]}
        for libelle, mots in cibles_par_menage.items():
            nom_cible = next((n for n in tables if _table_correspond(n, mots)), None)
            if nom_cible is None:
                continue
            resultat = controle_eligibilite_croisee_par_menage(tables, nom_presence, nom_cible)
            if resultat is not None:
                controles_croises.append((
                    f"Éligibilité présence ↔ {libelle} (par ménage)", nom_presence, nom_cible, resultat
                ))

    nom_deces = next((n for n in tables if _table_correspond(n, ["death", "deces"])), None)
    if nom_presence is not None and nom_deces is not None:
        resultat = controle_deces_present(tables, nom_deces, nom_presence)
        if resultat is not None:
            controles_croises.append(("Décédé mais présent dans la fiche présence", nom_deces, nom_presence, resultat))
    if nom_deces is not None:
        nom_individus_deces = _table_individus(tables, exclure=nom_deces)
        resultat = controle_deces_avant_naissance(tables, nom_deces)
        if resultat is not None:
            controles_croises.append((
                "Date de décès antérieure à la naissance", nom_deces, nom_individus_deces or nom_deces, resultat
            ))

    nom_relationship = next((n for n in tables if _table_correspond(n, ["relationship", "relation"])), None)
    nom_histmat = next((n for n in tables if _table_correspond(n, ["histmat", "matrimon", "marietal"])), None)
    if nom_relationship is not None:
        resultats_union = controle_naissance_apres_union(tables, nom_relationship)
        if resultats_union is not None:
            for libelle_union, resultat_union in resultats_union:
                controles_croises.append((libelle_union, nom_relationship, nom_relationship, resultat_union))
    if nom_histmat is not None and nom_relationship is not None:
        resultat = controle_ecart_age_union_declare_calcule(tables, nom_histmat, nom_relationship)
        if resultat is not None:
            controles_croises.append((
                "Âge à l'union déclaré très différent de l'âge calculé", nom_histmat, nom_relationship, resultat
            ))

    # "migration_out"/"migrationout" : anciens mots-cles synthetiques ;
    # "depart" : vrai nom reel le plus proche (table "opo_hypervel_departs"
    # - l'observatoire ne distingue pas de table "migration_out" a part,
    # un depart du menage est l'equivalent reel disponible).
    nom_migration_out = next(
        (n for n in tables if _table_correspond(n, ["migration_out", "migrationout", "depart"])), None
    )
    if nom_presence is not None and nom_migration_out is not None:
        resultat = controle_deces_present(tables, nom_migration_out, nom_presence)
        if resultat is not None:
            controles_croises.append(
                ("A dormi dans le ménage mais apparaît aussi en départ", nom_migration_out, nom_presence, resultat)
            )

    nom_migration_in = next((n for n in tables if _table_correspond(n, ["migration_in", "migrationin"])), None)
    if nom_migration_in is not None and nom_migration_out is not None:
        resultat = controle_migration_depart_avant_arrivee(tables, nom_migration_in, nom_migration_out)
        if resultat is not None:
            controles_croises.append((
                "Date de départ antérieure à la date d'arrivée (migration)",
                nom_migration_out, nom_migration_in, resultat,
            ))

    # "pregnancy"/"pregoutcome" : anciens mots-cles synthetiques ; "grossesse"
    # / "issue" : vrais noms reels (tables "opo_hypervel_grossesses" et
    # "opo_hypervel_issue_grossesses" - la deuxieme doit etre exclue de la
    # detection de la premiere, sous peine de se cibler elle-meme).
    nom_grossesse = next(
        (
            n
            for n in tables
            if _table_correspond(n, ["pregnancy", "grossesse"])
            and not _table_correspond(n, ["cpn", "outcome", "issue"])
        ),
        None,
    )
    nom_issue = next(
        (n for n in tables if _table_correspond(n, ["pregoutcome", "issuegrossesse", "issue_grossesse"])), None
    )
    if nom_grossesse is not None and nom_issue is not None:
        cle = detecter_cle_jointure(nom_grossesse, nom_issue, tables)
        if cle:
            manquantes = difference_tables(nom_grossesse, nom_issue, tables, cle=cle)
            controles_croises.append((
                "Grossesse sans issue de grossesse enregistrée", nom_grossesse, nom_issue,
                {"colonnes_verifiees": [cle], "n_anomalies": len(manquantes)},
            ))

    nom_naissance = next((n for n in tables if _table_correspond(n, ["birth", "naissance"])), None)
    if nom_naissance is not None and nom_issue is not None:
        resultat = controle_naissance_sans_issue(tables, nom_naissance, nom_issue)
        if resultat is not None:
            controles_croises.append((
                "Naissance sans issue de grossesse correspondante", nom_naissance, nom_issue, resultat
            ))

    if nom_issue is not None and nom_grossesse is not None:
        resultat = controle_issue_sans_grossesse(tables, nom_issue, nom_grossesse)
        if resultat is not None:
            controles_croises.append((
                "Issue de grossesse sans grossesse correspondante", nom_issue, nom_grossesse, resultat
            ))

    if nom_grossesse is not None and nom_presence is not None:
        resultat = controle_grossesse_sans_avoir_dormi(tables, nom_grossesse, nom_presence)
        if resultat is not None:
            controles_croises.append((
                "Enceinte sans avoir dormi sur place", nom_grossesse, nom_presence, resultat
            ))

    # "genesiq"/"gnesiq" EXCLUANT le complement - meme principe que
    # grossesse/issue plus haut (la table complementaire contient aussi
    # "genesiq" comme sous-chaine et se ciblerait sinon elle-meme).
    nom_genesique_base = next(
        (
            n for n in tables
            if _table_correspond(n, ["genesiq", "gnesiq"])
            and not _table_correspond(n, ["gnesiqcompl", "genesiqcompl", "genesiquecomplementaire"])
        ),
        None,
    )
    nom_complement = next(
        (n for n in tables if _table_correspond(n, ["gnesiqcompl", "genesiqcompl", "genesiquecomplementaire"])), None
    )
    for nom_cible_genesique in (nom_genesique_base, nom_complement):
        if nom_cible_genesique is None:
            continue
        resultat = controle_homme_dans_fiche_genesique(tables, nom_cible_genesique)
        if resultat is not None:
            controles_croises.append((
                "Homme présent dans la fiche génésique", nom_cible_genesique, nom_cible_genesique, resultat
            ))
    if nom_genesique_base is not None and nom_complement is not None:
        resultat = controle_pas_enfant_mais_naissance_vivante(tables, nom_genesique_base, nom_complement)
        if resultat is not None:
            controles_croises.append((
                "Aucun enfant déclaré alors qu'une naissance vivante est enregistrée",
                nom_genesique_base, nom_complement, resultat,
            ))
    if nom_complement is not None:
        resultat = controle_age_mere_a_naissance(tables, nom_complement)
        if resultat is not None:
            nom_individus_mere = _table_individus(tables, exclure=nom_complement)
            controles_croises.append((
                "Âge de la mère incohérent à la naissance (< 15 ans)",
                nom_complement, nom_individus_mere or nom_complement, resultat,
            ))

    nom_sante = next((n for n in tables if _table_correspond(n, ["sante"])), None)
    if nom_sante is not None and nom_presence is not None:
        resultat = controle_sante_repondant_non_dormi(tables, nom_sante, nom_presence)
        if resultat is not None:
            controles_croises.append(("Répondant santé n'ayant pas dormi sur place", nom_sante, nom_presence, resultat))
    if nom_sante is not None:
        nom_individus_sante = _table_individus(tables, exclure=nom_sante)
        resultat = controle_sante_repondant_mineur(tables, nom_sante)
        if resultat is not None:
            controles_croises.append((
                "Répondant santé mineur (< 15 ans)", nom_sante, nom_individus_sante or nom_sante, resultat
            ))
        resultat = controle_sante_regles_femme_jeune(tables, nom_sante)
        if resultat is not None:
            controles_croises.append((
                "Question sur les règles renseignée pour une femme de moins de 35 ans",
                nom_sante, nom_individus_sante or nom_sante, resultat,
            ))

    # "snakebite" seul (pas "snakebiteanimaux"/"snakebiteindividus", qui
    # portent des questions differentes, sans lien direct avec la residence
    # du menage).
    nom_snakebite = next(
        (
            n for n in tables
            if _table_correspond(n, ["snakebite"])
            and not _table_correspond(n, ["snakebiteanimaux", "snakebiteindividus"])
        ),
        None,
    )
    nom_residency = next((n for n in tables if _table_correspond(n, ["residency", "residence"])), None)
    if nom_snakebite is not None and nom_residency is not None:
        resultat = controle_snakebite_residents(tables, nom_snakebite, nom_residency)
        if resultat is not None:
            controles_croises.append((
                "Fiche snakebite ↔ résidents du ménage", nom_snakebite, nom_residency, resultat
            ))

    nom_telephone = next((n for n in tables if _table_correspond(n, ["telephone"])), None)
    if nom_telephone is not None:
        resultat = controle_telephone_par_age(tables, nom_telephone)
        if resultat is not None:
            nom_individus_tel = _table_individus(tables, exclure=nom_telephone)
            controles_croises.append((
                "Téléphone incohérent avec l'âge (< 15 ans avec numéro / ≥ 15 ans sans numéro)",
                nom_telephone, nom_individus_tel or nom_telephone, resultat,
            ))

    if nom_table is not None:
        controles_croises = [
            c for c in controles_croises if nom_table in (c[1], c[2])
        ]

    return {"par_table": resultats_par_table, "croises": controles_croises}


def controle_eligibilite_croisee(tables: dict, nom_presence: str, nom_cible: str, cle: str | None = None) -> dict | None:
    """Compare la population 'eligible' de la fiche presence (a dormi, sans
    date de depart enregistree) a une autre fiche (education/emploi/...) :
    qui a la fiche sans etre eligible, et qui est eligible sans avoir la
    fiche - repond directement a un des controles les plus demandes du
    catalogue de l'observatoire."""
    if nom_presence not in tables or nom_cible not in tables:
        return None
    presence, cible = tables[nom_presence], tables[nom_cible]
    cle = cle or detecter_cle_jointure(nom_presence, nom_cible, tables)
    if cle is None or cle not in presence.columns or cle not in cible.columns:
        return None

    col_sleep = _premiere_colonne(presence, SLEEP_LIKE)
    col_depart = _premiere_colonne(presence, DEPART_DATE_LIKE)
    masque = pd.Series(True, index=presence.index)
    colonnes_verifiees = [cle]
    if col_sleep is not None:
        masque &= presence[col_sleep].astype(str).str.strip().str.lower().isin(["1", "1.0", "oui", "yes", "true"])
        colonnes_verifiees.append(col_sleep)
    if col_depart is not None:
        masque &= presence[col_depart].isna()
        colonnes_verifiees.append(col_depart)

    ids_eligibles = set(presence.loc[masque, cle].dropna())
    ids_cible = set(cible[cle].dropna())
    eligibles_sans_fiche = sorted(ids_eligibles - ids_cible, key=str)
    fiche_sans_eligibilite = sorted(ids_cible - ids_eligibles, key=str)

    return {
        "colonnes_verifiees": colonnes_verifiees,
        "n_eligibles_sans_fiche": len(eligibles_sans_fiche),
        "n_fiche_sans_eligibilite": len(fiche_sans_eligibilite),
        "eligibles_sans_fiche": eligibles_sans_fiche[:50],
        "fiche_sans_eligibilite": fiche_sans_eligibilite[:50],
    }


def controle_deces_present(tables: dict, nom_source: str, nom_presence: str, cle: str | None = None) -> dict | None:
    """Compare les individus d'une table (deces, migration OUT...) a la fiche
    presence : signale ceux qui apparaissent dans les deux alors qu'ils ne
    devraient logiquement pas (ex: decede mais toujours marque present,
    parti en migration mais toujours marque comme ayant dormi sur place)."""
    if nom_source not in tables or nom_presence not in tables:
        return None
    cle = cle or detecter_cle_jointure(nom_source, nom_presence, tables)
    if cle is None:
        return None
    try:
        chevauchement = tables[nom_source].merge(tables[nom_presence][[cle]].drop_duplicates(), on=cle, how="inner")
    except Exception:
        return None
    return {"colonnes_verifiees": [cle], "n_anomalies": len(chevauchement)}


def controle_age_union(df: pd.DataFrame, borne_min: float = 12.0, borne_max: float = 40.0) -> dict | None:
    """Fiche histoire matrimoniale (FNewBase_HistMat) : signale les lignes ou
    l'age a la premiere union DECLARE (colonne `union_age`) sort de la
    tranche attendue. Controle DIRECT sur la valeur declaree - ne PAS
    confondre avec un age calcule depuis birth_date (qui donnerait l'age
    ACTUEL de la personne, pas son age au moment de l'union, tres different
    pour quelqu'un marie il y a des annees)."""
    col = _premiere_colonne(df, UNION_AGE_LIKE)
    if col is None:
        return None
    age = pd.to_numeric(df[col], errors="coerce")
    hors_plage = ((age < borne_min) | (age > borne_max)).fillna(False)
    return {
        "colonnes_verifiees": [col],
        "n_anomalies": int(hors_plage.sum()),
        "detail": [f"{col} hors {borne_min:g}-{borne_max:g} ans"],
    }


def controle_doublon_individu(df: pd.DataFrame) -> dict | None:
    """Signale un identifiant d'individu qui apparait PLUSIEURS FOIS dans une
    table ou il ne devrait logiquement apparaitre qu'une seule fois (ex: un
    meme individu declare decede a plus d'une reprise dans FNewDeath).

    A n'utiliser QUE pour les tables ou une repetition n'est PAS legitime -
    contrairement a une fiche d'evenement (grossesses, naissances...) ou un
    individu peut tres bien apparaitre plusieurs fois au fil du temps."""
    col_id = next((c for c in detect_id_columns(df) if "individ" in c.lower() and not c.lower().endswith("2")), None)
    if col_id is None:
        return None
    valeurs = df[col_id][df[col_id].notna()]
    doublons = valeurs[valeurs.duplicated(keep=False)]
    return {"colonnes_verifiees": [col_id], "n_anomalies": int(doublons.nunique())}


def controle_deces_formation_sanitaire(df: pd.DataFrame) -> dict | None:
    """FNewDeath : signale les deces declares survenus EN formation sanitaire
    (`diedinhs` = oui) alors que l'individu n'y est JAMAIS alle (`gonetohs` =
    non) - incoherent, un deces en formation sanitaire suppose necessairement
    un passage par cette formation avant le deces."""
    col_died = _premiere_colonne(df, DIEDINHS_LIKE)
    col_gone = _premiere_colonne(df, GONETOHS_LIKE)
    if col_died is None or col_gone is None:
        return None
    oui = lambda s: s.astype(str).str.strip().str.lower().isin(["1", "1.0", "oui", "yes", "true"])
    died_oui = oui(df[col_died])
    gone_non = df[col_gone].notna() & ~oui(df[col_gone])
    masque = died_oui & gone_non
    return {"colonnes_verifiees": [col_died, col_gone], "n_anomalies": int(masque.sum())}


def controle_deces_avant_naissance(tables: dict, nom_deces: str) -> dict | None:
    """FNewDeath : signale les deces dont la date (`evtdate`) est ANTERIEURE
    a la date de naissance de l'individu concerne, via `individid` -> table
    individus (voir `_table_individus`) - logiquement impossible."""
    if nom_deces not in tables:
        return None
    deces = tables[nom_deces]
    col_id = next((c for c in detect_id_columns(deces) if "individ" in c.lower() and not c.lower().endswith("2")), None)
    col_deces_date = _premiere_colonne(deces, DEATH_DATE_LIKE)
    if col_id is None or col_deces_date is None:
        return None
    nom_individus = _table_individus(tables, exclure=nom_deces)
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_naissance = _premiere_colonne(individus, BIRTH_DATE_LIKE)
    if col_id_individus is None or col_naissance is None:
        return None

    reference = individus[[col_id_individus, col_naissance]].rename(columns={col_id_individus: col_id})
    reference = reference.assign(**{col_id: _cle_str(reference[col_id])})
    gauche = deces[[col_id, col_deces_date]].assign(**{col_id: _cle_str(deces[col_id])})
    fusion = gauche.merge(reference, on=col_id, how="left")
    d_deces = pd.to_datetime(fusion[col_deces_date], errors="coerce", dayfirst=True)
    d_naissance = pd.to_datetime(fusion[col_naissance], errors="coerce", dayfirst=True)
    masque = (d_deces < d_naissance).fillna(False)
    return {
        "colonnes_verifiees": [col_id, col_deces_date, f"{nom_individus}.{col_naissance}"],
        "n_anomalies": int(masque.sum()),
    }


def controle_chef_menage_mineur(df: pd.DataFrame, seuil_age: float = 12.0) -> dict | None:
    """FNewBirth (ou toute fiche portant a la fois `birth_date` et
    `rltn_head`) : signale les individus de moins de `seuil_age` ans
    declares chef de menage (`rltn_head` = 1, "Socialgp head" - code
    documente dans le dictionnaire de donnees de l'observatoire, jamais
    devine) - logiquement improbable."""
    col_naissance = _premiere_colonne(df, BIRTH_DATE_LIKE)
    col_rltn = _premiere_colonne(df, RLTN_HEAD_LIKE)
    if col_naissance is None or col_rltn is None:
        return None
    age = _age_en_annees(df, col_naissance)
    est_chef = df[col_rltn].astype(str).str.strip() == CODE_CHEF_MENAGE
    masque = (age < seuil_age) & est_chef
    return {"colonnes_verifiees": [col_naissance, col_rltn], "n_anomalies": int(masque.sum())}


def controle_naissance_sans_issue(tables: dict, nom_naissance: str, nom_issue: str) -> dict | None:
    """FNewBirth <-> FNewPregoutcome : signale les naissances (`pregoutid`)
    sans AUCUNE correspondance dans la table des issues de grossesse
    (`eventid`) - cle PRECISE documentee par l'observatoire, differente d'une
    simple jointure par individid (qui donnerait un lien "large", pas la
    correspondance exacte demandee ici)."""
    if nom_naissance not in tables or nom_issue not in tables:
        return None
    naissance, issue = tables[nom_naissance], tables[nom_issue]
    col_pregoutid = _premiere_colonne(naissance, PREGOUTID_LIKE)
    col_eventid = _premiere_colonne(issue, EVENTID_LIKE)
    if col_pregoutid is None or col_eventid is None:
        return None
    ids_naissance = set(_cle_str(naissance[col_pregoutid]).dropna())
    ids_issue = set(_cle_str(issue[col_eventid]).dropna())
    manquants = ids_naissance - ids_issue
    return {
        "colonnes_verifiees": [f"{nom_naissance}.{col_pregoutid}", f"{nom_issue}.{col_eventid}"],
        "n_anomalies": len(manquants),
    }


def controle_issue_sans_grossesse(tables: dict, nom_issue: str, nom_grossesse: str) -> dict | None:
    """FNewPregoutcome <-> FNewPregnancy : signale les issues de grossesse
    (`peventid`) sans AUCUNE correspondance dans la table des grossesses
    (`eventid`) - cle PRECISE documentee par l'observatoire, dans le sens
    INVERSE de `controle_naissance_sans_issue`/du controle "grossesse sans
    issue" deja existant (qui, lui, utilise une cle plus large - individid)."""
    if nom_issue not in tables or nom_grossesse not in tables:
        return None
    issue, grossesse = tables[nom_issue], tables[nom_grossesse]
    col_peventid = _premiere_colonne(issue, PEVENTID_LIKE)
    col_eventid = _premiere_colonne(grossesse, EVENTID_LIKE)
    if col_peventid is None or col_eventid is None:
        return None
    ids_issue = set(_cle_str(issue[col_peventid]).dropna())
    ids_grossesse = set(_cle_str(grossesse[col_eventid]).dropna())
    manquants = ids_issue - ids_grossesse
    return {
        "colonnes_verifiees": [f"{nom_issue}.{col_peventid}", f"{nom_grossesse}.{col_eventid}"],
        "n_anomalies": len(manquants),
    }


def controle_grossesse_sans_avoir_dormi(tables: dict, nom_grossesse: str, nom_presence: str) -> dict | None:
    """FNewPregnancy <-> FNewPresences : signale les grossesses enregistrees
    pour un individu qui, selon la fiche presence, N'A PAS dormi sur place
    (`sleep_lastnight` different de oui) - incoherent."""
    if nom_grossesse not in tables or nom_presence not in tables:
        return None
    grossesse, presence = tables[nom_grossesse], tables[nom_presence]
    cle = detecter_cle_jointure(nom_grossesse, nom_presence, tables)
    if cle is None:
        return None
    col_sleep = _premiere_colonne(presence, SLEEP_LIKE)
    if col_sleep is None:
        return None
    dormi = presence[col_sleep].astype(str).str.strip().str.lower().isin(["1", "1.0", "oui", "yes", "true"])
    ids_non_dormi = set(_cle_str(presence.loc[~dormi, cle]).dropna())
    ids_grossesse = set(_cle_str(grossesse[cle]).dropna())
    return {
        "colonnes_verifiees": [cle, col_sleep],
        "n_anomalies": len(ids_grossesse & ids_non_dormi),
    }


def _colonnes_cpn_ordonnees(df: pd.DataFrame) -> list[str]:
    """Colonnes cpn_date1..cpn_dateN presentes dans la table, triees par
    NUMERO d'ordre (et non alphabetiquement, ce qui casserait des la
    dixieme visite : "cpn_date10" < "cpn_date2" en tri texte)."""
    trouvees = []
    for c in df.columns:
        m = CPN_DATE_LIKE.match(str(c))
        if m:
            trouvees.append((int(m.group(1)), c))
    return [c for _, c in sorted(trouvees)]


def controle_cpn_dates_desordonnees(df: pd.DataFrame) -> dict | None:
    """FNewPregnancy_CPN : signale les lignes ou les dates de visite CPN
    (cpn_date1, cpn_date2...) ne sont PAS dans l'ordre chronologique
    croissant (une visite numerotee plus tard datee avant une visite
    numerotee plus tot)."""
    colonnes = _colonnes_cpn_ordonnees(df)
    if len(colonnes) < 2:
        return None
    dates = [pd.to_datetime(df[c], errors="coerce", dayfirst=True) for c in colonnes]
    masque = pd.Series(False, index=df.index)
    for i in range(len(dates) - 1):
        for j in range(i + 1, len(dates)):
            masque |= (dates[i] > dates[j]).fillna(False)
    return {"colonnes_verifiees": colonnes, "n_anomalies": int(masque.sum())}


def controle_cpn_dates_manquantes(df: pd.DataFrame) -> dict | None:
    """FNewPregnancy_CPN : signale les lignes ou le nombre de CPN declare
    (`nb_cpn`) est SUPERIEUR au nombre de dates cpn_dateN reellement
    renseignees - des visites declarees sans date correspondante."""
    col_nb = _premiere_colonne(df, NB_CPN_LIKE)
    colonnes = _colonnes_cpn_ordonnees(df)
    if col_nb is None or not colonnes:
        return None
    nb_declare = pd.to_numeric(df[col_nb], errors="coerce")
    nb_renseignees = df[colonnes].notna().sum(axis=1)
    masque = (nb_declare > nb_renseignees).fillna(False)
    return {"colonnes_verifiees": [col_nb] + colonnes, "n_anomalies": int(masque.sum())}


def controle_homme_dans_fiche_genesique(tables: dict, nom_genesique: str) -> dict | None:
    """Fiche genesique (FNewHistoireGnesique/FNewHistGnesiqComplement) : ne
    devrait concerner QUE des femmes - signale les individus de genre
    masculin (`gender` = 1, "Male", code documente dans le dictionnaire de
    donnees de l'observatoire) presents dans la fiche, via individid -> table
    individus."""
    if nom_genesique not in tables:
        return None
    genesique = tables[nom_genesique]
    col_id = next(
        (c for c in detect_id_columns(genesique) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    if col_id is None:
        return None
    nom_individus = _table_individus(tables, exclure=nom_genesique)
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_gender = _premiere_colonne(individus, GENDER_LIKE)
    if col_id_individus is None or col_gender is None:
        return None
    reference = individus[[col_id_individus, col_gender]].rename(columns={col_id_individus: col_id})
    reference = reference.assign(**{col_id: _cle_str(reference[col_id])})
    gauche = genesique[[col_id]].assign(**{col_id: _cle_str(genesique[col_id])})
    fusion = gauche.merge(reference, on=col_id, how="left")
    masque = (fusion[col_gender].astype(str).str.strip() == CODE_HOMME).fillna(False)
    return {"colonnes_verifiees": [col_id, f"{nom_individus}.{col_gender}"], "n_anomalies": int(masque.sum())}


def controle_pas_enfant_mais_naissance_vivante(tables: dict, nom_genesique: str, nom_complement: str) -> dict | None:
    """FNewHistoireGnesique <-> FNewHistGnesiqComplement : signale les
    individus declarant AUCUN enfant vivant (`living_children_number` = 0)
    alors qu'une naissance vivante (`isAlive` = oui) est bien enregistree
    pour ce meme individu dans le complement genesique - incoherent."""
    if nom_genesique not in tables or nom_complement not in tables:
        return None
    genesique, complement = tables[nom_genesique], tables[nom_complement]
    col_id_g = next(
        (c for c in detect_id_columns(genesique) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_living = _premiere_colonne(genesique, LIVING_CHILDREN_LIKE)
    col_id_c = next(
        (c for c in detect_id_columns(complement) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_alive = _premiere_colonne(complement, ISALIVE_LIKE)
    if col_id_g is None or col_living is None or col_id_c is None or col_alive is None:
        return None
    nb_vivants = pd.to_numeric(genesique[col_living], errors="coerce")
    sans_enfant = set(_cle_str(genesique.loc[nb_vivants == 0, col_id_g]).dropna())
    oui = complement[col_alive].astype(str).str.strip().str.lower().isin(["1", "1.0", "oui", "yes", "true"])
    naissance_vivante = set(_cle_str(complement.loc[oui, col_id_c]).dropna())
    return {
        "colonnes_verifiees": [col_living, col_alive],
        "n_anomalies": len(sans_enfant & naissance_vivante),
    }


def controle_age_mere_a_naissance(tables: dict, nom_complement: str, seuil_age: float = 15.0) -> dict | None:
    """FNewHistGnesiqComplement : signale les cas ou l'age de la mere AU
    MOMENT de la naissance de l'enfant enregistre (`birthDate` de l'enfant -
    date de naissance de la mere, via individid -> table individus) est
    INFERIEUR a `seuil_age` ans - distinct d'un simple age ACTUEL hors
    tranche (voir `controle_tranche_age_croisee`), puisqu'il s'agit ici de
    l'age a un evenement passe."""
    if nom_complement not in tables:
        return None
    complement = tables[nom_complement]
    col_id = next(
        (c for c in detect_id_columns(complement) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_naissance_enfant = _premiere_colonne(complement, BIRTHDATE_ENFANT_LIKE)
    if col_id is None or col_naissance_enfant is None:
        return None
    nom_individus = _table_individus(tables, exclure=nom_complement)
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_naissance_mere = _premiere_colonne(individus, BIRTH_DATE_LIKE)
    if col_id_individus is None or col_naissance_mere is None:
        return None

    reference = individus[[col_id_individus, col_naissance_mere]].rename(
        columns={col_id_individus: col_id, col_naissance_mere: "_naissance_mere"}
    )
    reference = reference.assign(**{col_id: _cle_str(reference[col_id])})
    gauche = complement[[col_id, col_naissance_enfant]].assign(**{col_id: _cle_str(complement[col_id])})
    fusion = gauche.merge(reference, on=col_id, how="left")
    d_enfant = pd.to_datetime(fusion[col_naissance_enfant], errors="coerce", dayfirst=True)
    d_mere = pd.to_datetime(fusion["_naissance_mere"], errors="coerce", dayfirst=True)
    age_mere = (d_enfant - d_mere).dt.days / 365.25
    masque = (age_mere < seuil_age).fillna(False)
    return {
        "colonnes_verifiees": [col_id, col_naissance_enfant, f"{nom_individus}.{col_naissance_mere}"],
        "n_anomalies": int(masque.sum()),
    }


def controle_annee_naissance_enfant(df: pd.DataFrame, annee_min: int = 2022, annee_max: int = 2026) -> dict | None:
    """FNewHistGnesiqComplement : signale les enfants dont l'annee de
    naissance (`birthDate`) sort de la plage attendue des rounds de
    l'observatoire (2022-2026 par defaut)."""
    col = _premiere_colonne(df, BIRTHDATE_ENFANT_LIKE)
    if col is None:
        return None
    annee = pd.to_datetime(df[col], errors="coerce", dayfirst=True).dt.year
    masque = ((annee < annee_min) | (annee > annee_max)).fillna(False)
    return {
        "colonnes_verifiees": [col],
        "n_anomalies": int(masque.sum()),
        "detail": [f"{col} hors {annee_min}-{annee_max}"],
    }


def controle_telephone_par_age(tables: dict, nom_telephone: str, seuil_age: float = 15.0) -> dict | None:
    """FNewTelephone <-> table individus : signale (1) les individus de
    MOINS de `seuil_age` ans ayant un numero de telephone enregistre, et (2)
    les individus de `seuil_age` ans OU PLUS SANS aucun numero enregistre -
    a partir de l'age REEL (individid -> table individus), pas d'une simple
    presence brute dans FNewTelephone."""
    if nom_telephone not in tables:
        return None
    telephone = tables[nom_telephone]
    col_id_tel = next(
        (c for c in detect_id_columns(telephone) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    if col_id_tel is None:
        return None
    col_numero = _premiere_colonne(telephone, PHONE_LIKE)
    nom_individus = _table_individus(tables, exclure=nom_telephone)
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_naissance = _premiere_colonne(individus, BIRTH_DATE_LIKE)
    if col_id_individus is None or col_naissance is None:
        return None

    ids_avec_numero: set = set()
    if col_numero is not None:
        ids_avec_numero = set(_cle_str(telephone.loc[telephone[col_numero].notna(), col_id_tel]).dropna())

    reference = individus[[col_id_individus, col_naissance]].rename(columns={col_id_individus: "_id"})
    age = _age_en_annees(reference, col_naissance)
    reference = reference.assign(_age=age, _id=_cle_str(reference["_id"]))

    jeunes_avec_numero = reference.loc[(reference["_age"] < seuil_age) & reference["_id"].isin(ids_avec_numero)]
    ages_sans_numero = reference.loc[(reference["_age"] >= seuil_age) & ~reference["_id"].isin(ids_avec_numero)]

    return {
        "colonnes_verifiees": [col_id_tel, f"{nom_individus}.{col_naissance}"] + ([col_numero] if col_numero else []),
        "n_moins_seuil_avec_numero": int(len(jeunes_avec_numero)),
        "n_seuil_ou_plus_sans_numero": int(len(ages_sans_numero)),
    }


def controle_migration_depart_avant_arrivee(tables: dict, nom_migration_in: str, nom_migration_out: str) -> dict | None:
    """FNewMigration_IN <-> FNewMigration_Out : signale, pour un meme
    individu, un depart (`depart_date`, Migration_Out) ANTERIEUR a son
    arrivee (`arrive_date`, Migration_IN) - logiquement impossible (on ne
    peut pas partir d'un lieu avant d'y etre arrive)."""
    if nom_migration_in not in tables or nom_migration_out not in tables:
        return None
    entree, sortie = tables[nom_migration_in], tables[nom_migration_out]
    col_id_in = next(
        (c for c in detect_id_columns(entree) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_id_out = next(
        (c for c in detect_id_columns(sortie) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_arrivee = _premiere_colonne(entree, ARRIVE_DATE_LIKE)
    col_depart = _premiere_colonne(sortie, DEPART_DATE_LIKE)
    if col_id_in is None or col_id_out is None or col_arrivee is None or col_depart is None:
        return None

    gauche = sortie[[col_id_out, col_depart]].rename(columns={col_id_out: "_id"})
    gauche = gauche.assign(_id=_cle_str(gauche["_id"]))
    droite = entree[[col_id_in, col_arrivee]].rename(columns={col_id_in: "_id"})
    droite = droite.assign(_id=_cle_str(droite["_id"]))
    fusion = gauche.merge(droite, on="_id", how="inner")
    d_depart = pd.to_datetime(fusion[col_depart], errors="coerce", dayfirst=True)
    d_arrivee = pd.to_datetime(fusion[col_arrivee], errors="coerce", dayfirst=True)
    masque = (d_depart < d_arrivee).fillna(False)
    return {
        "colonnes_verifiees": [f"{nom_migration_out}.{col_depart}", f"{nom_migration_in}.{col_arrivee}"],
        "n_anomalies": int(masque.sum()),
    }


def controle_eligibilite_croisee_par_menage(tables: dict, nom_presence: str, nom_cible: str) -> dict | None:
    """Variante de `controle_eligibilite_croisee` pour les controles
    D'ELIGIBILITE PAR MENAGE (`socialgpid`) plutot que par individu - cas de
    la fiche pauvrete (et sante) : contrairement a l'individid, la fiche
    presence ne porte PAS directement `socialgpid`, donc une comparaison
    directe presence <-> cible ne trouverait aucune colonne commune. Il faut
    d'abord deriver le menage de chaque individu ELIGIBLE (a dormi, sans date
    de depart) via la table individus (`individid` -> `socialgpid`), PUIS
    comparer l'ensemble de menages obtenu a celui de la table cible."""
    if nom_presence not in tables or nom_cible not in tables:
        return None
    presence, cible = tables[nom_presence], tables[nom_cible]
    col_id_presence = next(
        (c for c in detect_id_columns(presence) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_socialgp_cible = _premiere_colonne(cible, SOCIALGP_LIKE)
    if col_id_presence is None or col_socialgp_cible is None:
        return None
    nom_individus = _table_individus(tables, exclure=nom_presence)
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_socialgp_individus = _premiere_colonne(individus, SOCIALGP_LIKE)
    if col_id_individus is None or col_socialgp_individus is None:
        return None

    col_sleep = _premiere_colonne(presence, SLEEP_LIKE)
    col_depart = _premiere_colonne(presence, DEPART_DATE_LIKE)
    masque = pd.Series(True, index=presence.index)
    colonnes_verifiees = [col_id_presence, col_socialgp_cible]
    if col_sleep is not None:
        masque &= presence[col_sleep].astype(str).str.strip().str.lower().isin(["1", "1.0", "oui", "yes", "true"])
        colonnes_verifiees.append(col_sleep)
    if col_depart is not None:
        masque &= presence[col_depart].isna()
        colonnes_verifiees.append(col_depart)

    ids_eligibles = set(_cle_str(presence.loc[masque, col_id_presence]).dropna())
    reference = individus[[col_id_individus, col_socialgp_individus]].rename(columns={col_id_individus: "_id"})
    reference = reference.assign(_id=_cle_str(reference["_id"]))
    menages_eligibles = set(reference.loc[reference["_id"].isin(ids_eligibles), col_socialgp_individus].dropna())

    ids_cible = set(cible[col_socialgp_cible].dropna())
    eligibles_sans_fiche = sorted(menages_eligibles - ids_cible, key=str)
    fiche_sans_eligibilite = sorted(ids_cible - menages_eligibles, key=str)

    return {
        "colonnes_verifiees": colonnes_verifiees + [f"{nom_individus}.{col_socialgp_individus}"],
        "n_eligibles_sans_fiche": len(eligibles_sans_fiche),
        "n_fiche_sans_eligibilite": len(fiche_sans_eligibilite),
        "eligibles_sans_fiche": eligibles_sans_fiche[:50],
        "fiche_sans_eligibilite": fiche_sans_eligibilite[:50],
    }


def _sante_avec_age(tables: dict, nom_sante: str) -> tuple[pd.DataFrame, str, str] | None:
    """Fusionne FNewSante avec la table individus via `respondid` ->
    `individid` (convention du schema reel de l'observatoire : le
    `respondid` d'une fiche sante EST l'`individid` de la personne
    interrogee) et ajoute une colonne d'age calcule `_age`. Renvoie
    (fusion, colonne_respondid, nom_table_individus), ou None si le
    croisement necessaire n'est pas disponible."""
    if nom_sante not in tables:
        return None
    sante = tables[nom_sante]
    col_respondid = _premiere_colonne(sante, RESPONDID_LIKE)
    if col_respondid is None:
        return None
    nom_individus = _table_individus(tables, exclure=nom_sante)
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_naissance = _premiere_colonne(individus, BIRTH_DATE_LIKE)
    if col_id_individus is None or col_naissance is None:
        return None
    reference = individus[[col_id_individus, col_naissance]].rename(columns={col_id_individus: col_respondid})
    reference = reference.assign(**{col_respondid: _cle_str(reference[col_respondid])})
    gauche = sante.assign(**{col_respondid: _cle_str(sante[col_respondid])})
    fusion = gauche.merge(reference, on=col_respondid, how="left")
    fusion = fusion.assign(_age=_age_en_annees(fusion, col_naissance))
    return fusion, col_respondid, nom_individus


def controle_sante_repondant_non_dormi(tables: dict, nom_sante: str, nom_presence: str) -> dict | None:
    """FNewSante <-> FNewPresences : signale les fiches sante remplies pour
    un repondant qui, selon la fiche presence, N'A PAS dormi sur place
    (convention : `respondid` (Sante) correspond a l'`individid` (Presence)
    de la meme personne)."""
    if nom_sante not in tables or nom_presence not in tables:
        return None
    sante, presence = tables[nom_sante], tables[nom_presence]
    col_respondid = _premiere_colonne(sante, RESPONDID_LIKE)
    col_id_presence = next(
        (c for c in detect_id_columns(presence) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_sleep = _premiere_colonne(presence, SLEEP_LIKE)
    if col_respondid is None or col_id_presence is None or col_sleep is None:
        return None
    dormi = presence[col_sleep].astype(str).str.strip().str.lower().isin(["1", "1.0", "oui", "yes", "true"])
    ids_non_dormi = set(_cle_str(presence.loc[~dormi, col_id_presence]).dropna())
    ids_sante = set(_cle_str(sante[col_respondid]).dropna())
    return {"colonnes_verifiees": [col_respondid, col_sleep], "n_anomalies": len(ids_sante & ids_non_dormi)}


def controle_sante_doublon_menage(df: pd.DataFrame) -> dict | None:
    """FNewSante : signale un meme `respondid` associe a PLUSIEURS
    `socialgpid` differents - incoherent, un repondant appartient a un seul
    menage."""
    col_respondid = _premiere_colonne(df, RESPONDID_LIKE)
    col_socialgp = _premiere_colonne(df, SOCIALGP_LIKE)
    if col_respondid is None or col_socialgp is None:
        return None
    nb_menages = df.groupby(col_respondid)[col_socialgp].nunique()
    return {"colonnes_verifiees": [col_respondid, col_socialgp], "n_anomalies": int((nb_menages > 1).sum())}


def controle_sante_mois_regle_inconnu(df: pd.DataFrame) -> dict | None:
    """FNewSante : quand la repondante n'a PAS eu ses regles au cours des 12
    derniers mois (`S4_2` = 2, "Non" - code documente dans le questionnaire
    source), le mois de la derniere fois (`S4_2mm`) devrait etre renseigne -
    signale les cas ou il manque."""
    col_s4_2 = _premiere_colonne(df, S4_2_LIKE)
    col_mm = _premiere_colonne(df, S4_2MM_LIKE)
    if col_s4_2 is None or col_mm is None:
        return None
    non = df[col_s4_2].astype(str).str.strip() == CODE_NON
    masque = non & df[col_mm].isna()
    return {"colonnes_verifiees": [col_s4_2, col_mm], "n_anomalies": int(masque.sum())}


def controle_sante_internet_contradictoire(df: pd.DataFrame) -> dict | None:
    """FNewSante : signale les repondants declarant n'avoir JAMAIS utilise
    internet en general (`S5_1` = 7, "Jamais" - questionnaire source) alors
    qu'ils declarent une utilisation reelle (reponse 1 a 4, PAS "Jamais"/"Ne
    sait pas") dans au moins un lieu precis (`S5_2A` maison, `S5_2B` travail/
    ecole, `S5_2C` espace public) - incoherent."""
    col_s5_1 = _premiere_colonne(df, S5_1_LIKE)
    colonnes_lieux = [c for c in df.columns if S5_2_LIEU_LIKE.match(str(c))]
    if col_s5_1 is None or not colonnes_lieux:
        return None
    jamais_general = df[col_s5_1].astype(str).str.strip() == "7"
    utilise_quelque_part = pd.Series(False, index=df.index)
    for c in colonnes_lieux:
        utilise_quelque_part |= df[c].astype(str).str.strip().isin(["1", "2", "3", "4"])
    masque = jamais_general & utilise_quelque_part
    return {"colonnes_verifiees": [col_s5_1] + colonnes_lieux, "n_anomalies": int(masque.sum())}


def controle_sante_repondant_mineur(tables: dict, nom_sante: str, seuil_age: float = 15.0) -> dict | None:
    """FNewSante : le module sante ne concerne QUE les repondants de 15 ans
    et plus (documente explicitement dans le questionnaire source, "Module
    sante (répondant de 15 ans et plus)") - signale les fiches remplies pour
    un repondant plus jeune."""
    resultat = _sante_avec_age(tables, nom_sante)
    if resultat is None:
        return None
    fusion, col_respondid, nom_individus = resultat
    masque = (fusion["_age"] < seuil_age).fillna(False)
    return {"colonnes_verifiees": [col_respondid, f"{nom_individus}.birth_date"], "n_anomalies": int(masque.sum())}


def controle_sante_regles_femme_jeune(tables: dict, nom_sante: str, seuil_age: float = 35.0) -> dict | None:
    """FNewSante : la question sur les regles des 12 derniers mois (`S4_2`)
    n'est documentee, dans le questionnaire source, que "Pour les femmes
    agees de 35 ans et plus" - signale les fiches ou `S4_2` est renseignee
    pour un repondant de moins de `seuil_age` ans."""
    resultat = _sante_avec_age(tables, nom_sante)
    if resultat is None:
        return None
    fusion, col_respondid, nom_individus = resultat
    col_s4_2 = _premiere_colonne(fusion, S4_2_LIKE)
    if col_s4_2 is None:
        return None
    masque = (fusion["_age"] < seuil_age).fillna(False) & fusion[col_s4_2].notna()
    return {
        "colonnes_verifiees": [col_respondid, col_s4_2, f"{nom_individus}.birth_date"],
        "n_anomalies": int(masque.sum()),
    }


def controle_snakebite_residents(tables: dict, nom_snakebite: str, nom_residency: str) -> dict | None:
    """FNewSnakebite <-> FNewResidency (`res_status` = 1, "Resident" - code
    documente) : signale, par MENAGE (`socialgpid`), les fiches snakebite
    sans AUCUN resident identifie, et separement les menages avec des
    residents mais SANS fiche snakebite."""
    if nom_snakebite not in tables or nom_residency not in tables:
        return None
    snakebite, residency = tables[nom_snakebite], tables[nom_residency]
    col_socialgp_snake = _premiere_colonne(snakebite, SOCIALGP_LIKE)
    col_socialgp_res = _premiere_colonne(residency, SOCIALGP_LIKE)
    col_res_status = _premiere_colonne(residency, RES_STATUS_LIKE)
    if col_socialgp_snake is None or col_socialgp_res is None or col_res_status is None:
        return None
    residents = residency[residency[col_res_status].astype(str).str.strip() == CODE_RESIDENT]
    menages_avec_residents = set(_cle_str(residents[col_socialgp_res]).dropna())
    menages_snakebite = set(_cle_str(snakebite[col_socialgp_snake]).dropna())
    fiche_sans_resident = sorted(menages_snakebite - menages_avec_residents, key=str)
    residents_sans_fiche = sorted(menages_avec_residents - menages_snakebite, key=str)
    return {
        "colonnes_verifiees": [col_socialgp_snake, col_socialgp_res, col_res_status],
        "n_fiche_sans_resident": len(fiche_sans_resident),
        "n_residents_sans_fiche": len(residents_sans_fiche),
        "fiche_sans_resident": fiche_sans_resident[:50],
        "residents_sans_fiche": residents_sans_fiche[:50],
    }


def controle_duree_entretien(df: pd.DataFrame, seuil_minutes: float = 3.0) -> dict | None:
    """FNewObservation : signale les entretiens de duree ANORMALEMENT COURTE
    (< `seuil_minutes` minutes, `end_time` - `begin_time`), et separement les
    entretiens a la duree INCOHERENTE (fin AVANT le debut)."""
    col_begin = _premiere_colonne(df, BEGIN_TIME_LIKE)
    col_end = _premiere_colonne(df, END_TIME_LIKE)
    if col_begin is None or col_end is None:
        return None
    d_begin = pd.to_datetime(df[col_begin], errors="coerce")
    d_end = pd.to_datetime(df[col_end], errors="coerce")
    duree_min = (d_end - d_begin).dt.total_seconds() / 60
    incoherente = (duree_min < 0).fillna(False)
    trop_courte = ((duree_min >= 0) & (duree_min < seuil_minutes)).fillna(False)
    return {
        "colonnes_verifiees": [col_begin, col_end],
        "n_duree_trop_courte": int(trop_courte.sum()),
        "n_duree_incoherente": int(incoherente.sum()),
    }


def controle_naissance_apres_union(tables: dict, nom_relationship: str) -> list[tuple[str, dict]] | None:
    """FNewRelationship : signale, pour chacune des 4 dates d'union
    documentees (debut, civile, religieuse, traditionnelle), les lignes ou la
    date de naissance de l'individu (via `individid` -> table individus) est
    POSTERIEURE a la date d'union - logiquement impossible (on ne peut pas se
    marier avant de naitre).

    Renvoie une liste de (libelle, resultat) - un element par variante de
    date effectivement presente dans la table - ou None si la table/le
    croisement necessaire n'est pas disponible."""
    if nom_relationship not in tables:
        return None
    relationship = tables[nom_relationship]
    col_id = next(
        (c for c in detect_id_columns(relationship) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    if col_id is None:
        return None
    nom_individus = _table_individus(tables, exclure=nom_relationship)
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_naissance = _premiere_colonne(individus, BIRTH_DATE_LIKE)
    if col_id_individus is None or col_naissance is None:
        return None

    reference = individus[[col_id_individus, col_naissance]].rename(columns={col_id_individus: col_id})
    naissance = pd.to_datetime(reference[col_naissance], errors="coerce", dayfirst=True)
    reference = reference.assign(**{col_naissance: naissance, col_id: _cle_str(reference[col_id])})

    resultats = []
    for libelle, motif in DATES_UNION_LIKE.items():
        col_union = _premiere_colonne(relationship, motif)
        if col_union is None:
            continue
        gauche = relationship[[col_id, col_union]].assign(**{col_id: _cle_str(relationship[col_id])})
        fusion = gauche.merge(reference, on=col_id, how="left")
        d_union = pd.to_datetime(fusion[col_union], errors="coerce", dayfirst=True)
        d_naissance = fusion[col_naissance]
        masque = (d_naissance > d_union).fillna(False)
        resultats.append((
            f"Naissance postérieure à la date d'union ({libelle})",
            {
                "colonnes_verifiees": [col_id, col_union, f"{nom_individus}.{col_naissance}"],
                "n_anomalies": int(masque.sum()),
            },
        ))
    return resultats or None


def controle_ecart_age_union_declare_calcule(
    tables: dict, nom_histmat: str, nom_relationship: str, seuil_ecart: float = 3.0
) -> dict | None:
    """FNewBase_HistMat <-> FNewRelationship <-> table individus : compare
    l'age a l'union DECLARE (`union_age`, FNewBase_HistMat) a l'age CALCULE
    (date d'union de debut - date de naissance) - signale un ecart de plus de
    `seuil_ecart` ans, indice d'une erreur de saisie sur l'une des deux
    dates/valeurs. Utilise la date d'union de DEBUT (`uni_start_date`) comme
    reference, faute d'un lien direct entre `union_age` et l'une des 4
    variantes de date documentees."""
    if nom_histmat not in tables or nom_relationship not in tables:
        return None
    histmat = tables[nom_histmat]
    relationship = tables[nom_relationship]
    col_id_histmat = next(
        (c for c in detect_id_columns(histmat) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_age_declare = _premiere_colonne(histmat, UNION_AGE_LIKE)
    if col_id_histmat is None or col_age_declare is None:
        return None

    col_id_relationship = next(
        (c for c in detect_id_columns(relationship) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_union = _premiere_colonne(relationship, DATES_UNION_LIKE["début d'union"])
    if col_id_relationship is None or col_union is None:
        return None

    nom_individus = _table_individus(tables, exclure=nom_histmat)
    if nom_individus is None or nom_individus == nom_relationship:
        nom_individus = _table_individus(tables) if nom_individus is None else nom_individus
    if nom_individus is None:
        return None
    individus = tables[nom_individus]
    col_id_individus = next(
        (c for c in detect_id_columns(individus) if "individ" in c.lower() and not c.lower().endswith("2")), None
    )
    col_naissance = _premiere_colonne(individus, BIRTH_DATE_LIKE)
    if col_id_individus is None or col_naissance is None:
        return None

    fusion = (
        histmat[[col_id_histmat, col_age_declare]]
        .rename(columns={col_id_histmat: "_id"})
        .assign(_id=lambda d: _cle_str(d["_id"]))
        .merge(
            relationship[[col_id_relationship, col_union]]
            .rename(columns={col_id_relationship: "_id"})
            .assign(_id=lambda d: _cle_str(d["_id"])),
            on="_id", how="inner",
        )
        .merge(
            individus[[col_id_individus, col_naissance]]
            .rename(columns={col_id_individus: "_id"})
            .assign(_id=lambda d: _cle_str(d["_id"])),
            on="_id", how="inner",
        )
    )
    if fusion.empty:
        return {
            "colonnes_verifiees": [col_id_histmat, col_age_declare, col_union, f"{nom_individus}.{col_naissance}"],
            "n_anomalies": 0,
        }

    age_declare = pd.to_numeric(fusion[col_age_declare], errors="coerce")
    d_union = pd.to_datetime(fusion[col_union], errors="coerce", dayfirst=True)
    d_naissance = pd.to_datetime(fusion[col_naissance], errors="coerce", dayfirst=True)
    age_calcule = (d_union - d_naissance).dt.days / 365.25
    ecart = (age_declare - age_calcule).abs()
    masque = (ecart > seuil_ecart).fillna(False)
    return {
        "colonnes_verifiees": [col_id_histmat, col_age_declare, col_union, f"{nom_individus}.{col_naissance}"],
        "n_anomalies": int(masque.sum()),
        "detail": [f"écart > {seuil_ecart:g} ans entre âge déclaré et âge calculé"],
    }


"""
Module "Performances" : suivi du volume d'activite de terrain par agent
enqueteur (menages/UCH visites, naissances/deces/grossesses enregistres),
independant du controle qualite deja fourni par `rapport_agents` (qui mesure
les erreurs/doublons, pas le volume). Meme principe d'auto-detection par nom
de colonne/table que le reste du module : rien n'est code en dur sur des
identifiants d'agents ou une equipe reelle, puisqu'ils varient d'un
observatoire/campagne a l'autre et ne sont pas connus a l'avance.
"""

CONTROLEUR_LIKE = re.compile(
    r"(controleur|contr[oô]lleur|superviseur|supervisor|chef_?equipe|team_?lead|^contro$)", re.IGNORECASE
)
# `^contro$` (colonne nommee EXACTEMENT "contro", rien avant/apres) : cas reel
# rencontre dans un export Stata (.dta) d'equipe agent<->controleur, ou le nom
# de colonne "Controleur" a ete tronque a "Contro" - ancre avec ^...$ plutot
# qu'un simple "contro" en sous-chaine pour ne jamais matcher par erreur une
# colonne de controle QUALITE (ex: "controle_qualite", "date_controle"), qui
# n'a rien a voir avec l'identite d'un superviseur.

# Chaque entree : (libelle affiche, mots-cles de reconnaissance du role de la
# table par son nom - meme principe que `_table_correspond`/`TRANCHES_AGE_PAR_TYPE_FICHE`,
# compter_menages -> si True, compte aussi les menages/UCH distincts (colonne
# LOCATION_LIKE) en plus du nombre de fiches.
CATEGORIES_PERFORMANCE_TERRAIN = [
    ("Ménages/UCH visités", ["presences", "presence"], True),
    ("Naissances enregistrées", ["birth", "naissance"], False),
    ("Décès enregistrés", ["death", "deces"], False),
    ("Grossesses enregistrées", ["pregnancy", "grossesse"], False),
]

# Colonne de cle etrangere vers l'enquete/visite de terrain a laquelle une
# fiche se rattache (ex: "enquete_id") - schema reel de l'observatoire (base
# Hypervel) : l'identite de l'agent n'est saisie qu'UNE FOIS, sur la table
# "opo_hypervel_enquete_or_visites", pas directement sur chaque fiche
# individuelle (naissances, deces, education...) qui ne porte que cette cle.
ENQUETE_ID_LIKE = re.compile(r"(enquete_?id|enquete_?or_?visite_?id|visite_?id)", re.IGNORECASE)

# Adresse email de l'agent, seule identite disponible dans la table
# "opo_hypervel_users" du schema reel (qui ne porte pas de nom/prenom).
EMAIL_LIKE = re.compile(r"(email|e_?mail|courriel)", re.IGNORECASE)


def _colonne_id_primaire(df: pd.DataFrame) -> str | None:
    """Colonne d'identifiant primaire d'une table (typiquement `id`) - a
    distinguer d'une simple cle etrangere detectee par `detect_id_columns`,
    car necessaire pour joindre une table de reference (enquetes/visites,
    utilisateurs) a une cle qui la cite ailleurs."""
    exact = next((c for c in df.columns if str(c).strip().lower() == "id"), None)
    if exact is not None:
        return exact
    return next(iter(detect_id_columns(df)), None)


def _table_enquetes(tables: dict) -> tuple[str, pd.DataFrame, str, str] | None:
    """Trouve la table de reference "enquetes/visites" : celle qui porte a la
    fois une colonne d'agent detectee directement ET un identifiant primaire
    - c'est elle qui permet de retrouver l'agent responsable d'une fiche qui
    ne cite qu'un `enquete_id` (voir `colonne_agent_effective`).

    Renvoie (nom_table, dataframe, colonne_id, colonne_agent) ou None si
    aucune table de ce type n'est chargee."""
    for nom, df in tables.items():
        if not _table_correspond(nom, ["enquete", "visite"]):
            continue
        col_agent = next(iter(detect_agent_columns(df)), None)
        col_id = _colonne_id_primaire(df)
        if col_agent is None or col_id is None:
            continue
        return nom, df, col_id, col_agent
    return None


def colonne_agent_effective(df: pd.DataFrame, tables: dict) -> tuple[pd.DataFrame, str | None]:
    """Determine quelle colonne utiliser pour identifier l'agent sur une
    fiche donnee, et enrichit au besoin le dataframe par jointure :

    1. Colonne agent directe (ex: `field_wrkr`, `agent_id`) si presente -
       cas le plus simple, notamment les donnees de test historiques.
    2. Sinon, si la fiche porte une cle `enquete_id` ET qu'une table
       enquetes/visites avec agent est chargee (voir `_table_enquetes`),
       l'agent est deduit par jointure et ajoute EN MEMOIRE (jamais ecrit
       dans les fichiers source) sous une colonne dediee
       `__agent_via_enquete__` - c'est le cas du schema reel de
       l'observatoire, ou l'identite de l'agent n'est saisie qu'une fois par
       enquete/visite, pas sur chaque fiche individuelle qui s'y rattache.

    Renvoie (df eventuellement enrichi, nom de colonne a utiliser, ou None si
    aucun agent n'a pu etre determine)."""
    col_direct = next(iter(detect_agent_columns(df)), None)
    if col_direct is not None:
        return df, col_direct

    col_enquete_id = _premiere_colonne(df, ENQUETE_ID_LIKE)
    if col_enquete_id is None:
        return df, None

    ref = _table_enquetes(tables)
    if ref is None:
        return df, None
    _, df_enquetes, col_id, col_agent_ref = ref
    if col_enquete_id not in df.columns:
        return df, None

    mapping = (
        df_enquetes[[col_id, col_agent_ref]]
        .drop_duplicates(subset=[col_id])
        .set_index(col_id)[col_agent_ref]
    )
    df = df.copy()
    df["__agent_via_enquete__"] = df[col_enquete_id].map(mapping)
    if df["__agent_via_enquete__"].notna().sum() == 0:
        return df, None
    return df, "__agent_via_enquete__"


def rapport_performance_agents(
    tables: dict, exclure: list[str] | None = None, colonne_agent: str | None = None
) -> pd.DataFrame:
    """Volume d'activite de terrain par agent, agrege a travers TOUTES les
    tables chargees qui comportent une colonne d'agent detectable (pas une
    seule table par defaut - meme principe que le reste de l'assistant) :
    nombre de menages/UCH visites, naissances/deces/grossesses enregistres,
    plus un total "autres fiches" pour toute table avec un agent qui ne
    correspond a aucune categorie reconnue (transparence : on ne perd aucune
    activite de terrain juste parce qu'elle ne rentre pas dans les 4
    categories nommees).

    `exclure` : liste d'identifiants d'agent (ex: agents non-terrain,
    formateurs, superviseurs saisis par erreur comme agent) a retirer du
    rapport - comparaison insensible a la casse/aux espaces, jamais une liste
    figee en dur puisque l'observatoire est seul a savoir qui exclure.
    """
    exclure_norm = {str(a).strip().lower() for a in exclure} if exclure else set()
    donnees: dict[str, dict[str, float]] = {}

    for nom, df in tables.items():
        if colonne_agent and colonne_agent in df.columns:
            col_agent = colonne_agent
        else:
            # Colonne agent directe si presente ; sinon, jointure automatique
            # via `enquete_id` vers la table enquetes/visites (schema reel de
            # l'observatoire ou l'identite de l'agent n'est saisie qu'une
            # fois par enquete, pas sur chaque fiche) - voir
            # `colonne_agent_effective`.
            df, col_agent = colonne_agent_effective(df, tables)
        if col_agent is None:
            continue
        # Une table qui comporte une colonne "controleur" est traitee comme la
        # table d'equipe (mapping agent <-> controleur, voir
        # `fusion_agent_controleur`) et non comme une fiche d'activite de
        # terrain : elle ne doit pas etre comptee ici, sous peine de gonfler
        # artificiellement le volume de chaque agent avec des lignes qui ne
        # sont pas des fiches saisies.
        if _premiere_colonne(df, CONTROLEUR_LIKE) is not None:
            continue

        categorie, compter_menages = next(
            (
                (libelle, cm)
                for libelle, mots, cm in CATEGORIES_PERFORMANCE_TERRAIN
                if _table_correspond(nom, mots)
            ),
            (f"Autres fiches ({nom})", False),
        )

        col_menage = _premiere_colonne(df, LOCATION_LIKE) if compter_menages else None
        for agent, index_groupe in df.groupby(col_agent, dropna=False).groups.items():
            agent_str = str(agent).strip()
            if agent_str.lower() in exclure_norm:
                continue
            entree = donnees.setdefault(agent_str, {})
            entree[categorie] = entree.get(categorie, 0) + len(index_groupe)
            if col_menage is not None:
                n_menages = int(df.loc[index_groupe, col_menage].nunique())
                cle_menages = "Ménages/UCH distincts"
                entree[cle_menages] = entree.get(cle_menages, 0) + n_menages

    if not donnees:
        return pd.DataFrame(columns=["agent", "total_fiches"])

    rapport = pd.DataFrame.from_dict(donnees, orient="index").fillna(0)
    rapport = rapport.reset_index().rename(columns={"index": "agent"})
    colonnes_valeurs = [c for c in rapport.columns if c != "agent"]
    for c in colonnes_valeurs:
        rapport[c] = rapport[c].astype(int)
    rapport["total_fiches"] = rapport[[c for c in colonnes_valeurs if c != "Ménages/UCH distincts"]].sum(axis=1)
    colonnes_ordre = ["agent"] + sorted(c for c in colonnes_valeurs) + ["total_fiches"]
    return rapport[colonnes_ordre].sort_values("total_fiches", ascending=False).reset_index(drop=True)


def rapport_performance_par_jour(tables: dict, colonne_agent: str | None = None) -> pd.DataFrame:
    """Volume de fiches par agent ET par jour (forme longue : date, agent,
    n_fiches), a partir de la premiere table qui comporte a la fois une
    colonne d'agent et une colonne de date detectees - en priorite la fiche
    presence (role le plus proche d'un "passage terrain" quotidien), sinon la
    premiere table eligible trouvee parmi celles chargees."""
    candidates = []
    for nom, df in tables.items():
        if colonne_agent and colonne_agent in df.columns:
            col_agent = colonne_agent
        else:
            # Meme jointure de repli via `enquete_id` que dans
            # `rapport_performance_agents` (voir `colonne_agent_effective`).
            df, col_agent = colonne_agent_effective(df, tables)
        if col_agent is None or _premiere_colonne(df, CONTROLEUR_LIKE) is not None:
            continue
        colonnes_date = detect_date_columns(df)
        if not colonnes_date:
            continue
        priorite = 0 if _table_correspond(nom, ["presences", "presence"]) else 1
        candidates.append((priorite, nom, df, col_agent, colonnes_date[0]))

    if not candidates:
        return pd.DataFrame(columns=["date", "agent", "n_fiches"])

    _, nom, df, col_agent, col_date = sorted(candidates, key=lambda c: c[0])[0]
    dates = pd.to_datetime(df[col_date], errors="coerce", dayfirst=True).dt.date
    temp = pd.DataFrame({"date": dates, "agent": df[col_agent].astype(str).str.strip()})
    temp = temp.dropna(subset=["date"])
    rapport = temp.groupby(["date", "agent"], dropna=False).size().reset_index(name="n_fiches")
    return rapport.sort_values(["date", "agent"]).reset_index(drop=True)


def fusion_agent_controleur(rapport_agents: pd.DataFrame, tables: dict) -> tuple[pd.DataFrame, str | None]:
    """Ajoute une colonne `controleur` au rapport de performance par agent, si
    une table "equipe" (contenant a la fois une colonne d'agent ET une
    colonne de controleur/superviseur, detectees par leur nom) est chargee -
    jointure agent <-> controleur demandee par l'observatoire pour situer
    chaque agent dans son equipe d'encadrement.

    Renvoie (rapport enrichi, nom de la table equipe utilisee) ; si aucune
    table de ce type n'est chargee, renvoie (rapport inchange, None) plutot
    que d'echouer."""
    if rapport_agents.empty or "agent" not in rapport_agents.columns:
        return rapport_agents, None

    for nom, df in tables.items():
        col_agent = next(iter(detect_agent_columns(df)), None)
        col_controleur = _premiere_colonne(df, CONTROLEUR_LIKE)
        if col_agent is None or col_controleur is None or col_agent == col_controleur:
            continue
        mapping = df[[col_agent, col_controleur]].drop_duplicates(subset=[col_agent])
        mapping = mapping.rename(columns={col_agent: "agent", col_controleur: "controleur"})
        mapping["agent"] = mapping["agent"].astype(str).str.strip()
        fusionne = rapport_agents.merge(mapping, on="agent", how="left")
        fusionne["controleur"] = fusionne["controleur"].fillna("Non renseigné")
        colonnes = ["agent", "controleur"] + [c for c in fusionne.columns if c not in ("agent", "controleur")]
        return fusionne[colonnes], nom

    return rapport_agents, None


def fusion_identite_agent(rapport_agents: pd.DataFrame, tables: dict) -> tuple[pd.DataFrame, str | None]:
    """Ajoute une colonne `email_agent` au rapport de performance par agent,
    en joignant l'identifiant agent a la table des utilisateurs (detectee par
    son nom - "users"/"utilisateurs"), si elle est chargee.

    C'est la SEULE identite disponible pour un agent dans le schema reel de
    l'observatoire : la table "opo_hypervel_users" ne porte pas de colonne
    nom/prenom (deliberement retiree cote source), uniquement un identifiant
    et une adresse email - c'est donc l'email qui permet de savoir "quel
    agent a fait quoi", pas un nom.

    Meme principe que `fusion_agent_controleur` : renvoie (rapport enrichi,
    nom de la table utilisateurs utilisee) ou (rapport inchange, None) si
    aucune table de ce type n'est chargee, plutot que d'echouer."""
    if rapport_agents.empty or "agent" not in rapport_agents.columns:
        return rapport_agents, None

    for nom, df in tables.items():
        if not _table_correspond(nom, ["users", "utilisateurs", "user"]):
            continue
        col_email = _premiere_colonne(df, EMAIL_LIKE)
        col_id = _colonne_id_primaire(df)
        if col_email is None or col_id is None:
            continue
        mapping = df[[col_id, col_email]].drop_duplicates(subset=[col_id])
        mapping = mapping.rename(columns={col_id: "agent", col_email: "email_agent"})
        mapping["agent"] = mapping["agent"].astype(str).str.strip()
        fusionne = rapport_agents.merge(mapping, on="agent", how="left")
        fusionne["email_agent"] = fusionne["email_agent"].fillna("Non renseigné")
        colonnes = ["agent", "email_agent"] + [c for c in fusionne.columns if c not in ("agent", "email_agent")]
        return fusionne[colonnes], nom

    return rapport_agents, None


def prevision_objectif(
    rapport_par_jour: pd.DataFrame, objectif: int = 17000, colonne_valeur: str = "n_fiches"
) -> dict | None:
    """Projection simple (rythme moyen constant) pour atteindre un objectif
    configurable (par defaut 17000 menages, chiffre cible communique par
    l'observatoire) a partir du cumul quotidien observe - renvoie le cumul
    actuel, le rythme journalier moyen, le nombre de jours restants estimes et
    la date de fin projetee. Renvoie None si aucune donnee par jour
    n'est disponible (rien a projeter)."""
    if rapport_par_jour is None or rapport_par_jour.empty:
        return None

    par_jour = rapport_par_jour.groupby("date")[colonne_valeur].sum().sort_index()
    cumul_actuel = int(par_jour.sum())
    n_jours = len(par_jour)
    rythme_journalier = cumul_actuel / n_jours if n_jours else 0.0
    restant = max(objectif - cumul_actuel, 0)

    resultat = {
        "objectif": objectif,
        "cumul_actuel": cumul_actuel,
        "n_jours_observes": n_jours,
        "rythme_journalier_moyen": round(rythme_journalier, 1),
        "reste_a_faire": restant,
        "date_debut": str(par_jour.index.min()),
        "date_derniere_donnee": str(par_jour.index.max()),
    }

    if restant == 0:
        resultat["jours_restants_estimes"] = 0
        resultat["date_fin_projetee"] = resultat["date_derniere_donnee"]
    elif rythme_journalier > 0:
        jours_restants = int(np.ceil(restant / rythme_journalier))
        resultat["jours_restants_estimes"] = jours_restants
        resultat["date_fin_projetee"] = str(
            (pd.Timestamp(resultat["date_derniere_donnee"]) + pd.Timedelta(days=jours_restants)).date()
        )
    else:
        resultat["jours_restants_estimes"] = None
        resultat["date_fin_projetee"] = None

    return resultat


def simulation_rythme(reste_a_faire: int, jours_disponibles: int) -> dict:
    """Simulation simple : rythme journalier necessaire pour finir le
    `reste_a_faire` (ex: menages restants) dans un nombre de jours donne -
    reponse a une question hypothetique ("si on veut finir en 20 jours, il
    faut combien par jour ?")."""
    if jours_disponibles <= 0:
        raise ValueError("Le nombre de jours disponibles doit être positif.")
    return {
        "reste_a_faire": reste_a_faire,
        "jours_disponibles": jours_disponibles,
        "rythme_journalier_necessaire": round(reste_a_faire / jours_disponibles, 1),
    }


def rechercher_identifiant(identifiant: str, tables: dict) -> dict[str, pd.DataFrame]:
    """Recherche instantanee d'un identifiant a travers TOUTES les tables
    chargees (pas seulement la table active) : renvoie, pour chaque table ou
    au moins une colonne d'identifiant contient cette valeur, les lignes
    correspondantes - pour retrouver en un coup toutes les fiches liees a un
    meme individu/menage sans avoir a nommer chaque table une par une."""
    identifiant_norm = str(identifiant).strip().lower()
    resultats: dict[str, pd.DataFrame] = {}
    for nom, df in tables.items():
        colonnes_id = detect_id_columns(df)
        if not colonnes_id:
            continue
        masque = pd.Series(False, index=df.index)
        for col in colonnes_id:
            masque |= df[col].astype(str).str.strip().str.lower() == identifiant_norm
        if masque.any():
            resultats[nom] = df[masque]
    return resultats


def generer_rapport_performance_docx(
    rapport_agents: pd.DataFrame, prevision: dict | None = None, objectif: int = 17000
) -> bytes:
    """Genere un rapport Word telechargeable (python-docx, deja une
    dependance du projet) resumant la performance de terrain par agent et,
    si disponible, la projection vers l'objectif configurable - pour partager
    un point d'avancement hors du chat (reunion d'equipe, hierarchie...)."""
    from docx import Document

    document = Document()
    document.add_heading("Rapport de performance de terrain", level=1)
    document.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}.")

    if prevision:
        document.add_heading("Avancement vers l'objectif", level=2)
        document.add_paragraph(
            f"Objectif : {prevision.get('objectif', objectif)} — "
            f"réalisé à ce jour : {prevision.get('cumul_actuel', 0)} — "
            f"reste à faire : {prevision.get('reste_a_faire', 0)}."
        )
        document.add_paragraph(
            f"Rythme journalier moyen observé : {prevision.get('rythme_journalier_moyen', 0)} / jour "
            f"sur {prevision.get('n_jours_observes', 0)} jour(s) de données "
            f"({prevision.get('date_debut', '?')} → {prevision.get('date_derniere_donnee', '?')})."
        )
        if prevision.get("date_fin_projetee"):
            document.add_paragraph(
                f"Date de fin projetée au rythme actuel : {prevision['date_fin_projetee']} "
                f"(≈ {prevision.get('jours_restants_estimes', '?')} jour(s) restants)."
            )

    document.add_heading("Performance par agent", level=2)
    if rapport_agents is None or rapport_agents.empty:
        document.add_paragraph("Aucune donnée de performance disponible.")
    else:
        colonnes = list(rapport_agents.columns)
        table = document.add_table(rows=1, cols=len(colonnes))
        table.style = "Light Grid Accent 1"
        for cell, nom_colonne in zip(table.rows[0].cells, colonnes):
            cell.text = str(nom_colonne)
        for _, ligne in rapport_agents.iterrows():
            cells = table.add_row().cells
            for cell, valeur in zip(cells, ligne):
                cell.text = str(valeur)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _noms_compatibles_stata(df: pd.DataFrame) -> pd.DataFrame:
    """Adapte les noms de colonnes aux contraintes Stata (<= 32 caracteres,
    commence par une lettre/underscore, uniquement lettres/chiffres/underscore),
    sans modifier le DataFrame original."""
    df = df.copy()
    nouveaux_noms = {}
    vus = set()
    for c in df.columns:
        nom = re.sub(r"[^A-Za-z0-9_]", "_", str(c))
        if not nom or not (nom[0].isalpha() or nom[0] == "_"):
            nom = "v_" + nom
        nom = nom[:32]
        base, i = nom, 1
        while nom in vus:
            suffix = f"_{i}"
            nom = base[: 32 - len(suffix)] + suffix
            i += 1
        vus.add(nom)
        nouveaux_noms[c] = nom
    return df.rename(columns=nouveaux_noms)


def exporter(df: pd.DataFrame, format: str) -> bytes:
    """Convertit une table en bytes pretes a telecharger, dans le format demande.

    format : "csv", "xlsx" ou "dta" (Stata).
    """
    format = format.lower().lstrip(".")
    buffer = io.BytesIO()

    if format == "csv":
        buffer.write(df.to_csv(index=False).encode("utf-8-sig"))

    elif format in ("xlsx", "excel"):
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="donnees")

    elif format in ("dta", "stata"):
        df_stata = _noms_compatibles_stata(df)
        # Stata n'accepte pas les colonnes entierement vides de type object,
        # ni les booleens : on les convertit en texte pour eviter une erreur.
        for c in df_stata.columns:
            if df_stata[c].dtype == bool:
                df_stata[c] = df_stata[c].astype(str)
        df_stata.to_stata(buffer, write_index=False, version=118)

    else:
        raise ValueError(f"Format d'export non supporte : {format} (attendu : csv, xlsx, dta)")

    buffer.seek(0)
    return buffer.getvalue()
